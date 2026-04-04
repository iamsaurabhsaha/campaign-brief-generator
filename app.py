"""
Campaign Brief Generator Dashboard
====================================
AI-powered Campaign Brief Generator that guides PMMs through writing
great campaign briefs using industry best practices. Interactive assistant
that helps users think through each section, catches common mistakes,
and produces professional output.
"""

import streamlit as st
import pandas as pd
import plotly.express as px
import plotly.graph_objects as go
from datetime import datetime, timedelta
from typing import Any, Optional
import json
import io
import os
import sys
import uuid
import logging
from pathlib import Path
from dotenv import load_dotenv

# Load .env from project root
load_dotenv(Path(__file__).parent / ".env")

# Add project root to path for local imports
sys.path.insert(0, str(Path(__file__).parent))

from brief_generator import BriefGenerator, BriefStore
from quality_checker import QualityChecker
from creative_engine import CreativeEngine
from templates import TemplateManager
from config import *

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Color scheme & constants
# ---------------------------------------------------------------------------
PRIMARY_BLUE = "#0064D2"
ACCENT_RED = "#E53238"
ACCENT_YELLOW = "#F5AF02"
ACCENT_GREEN = "#86B817"

TIER_DESCRIPTIONS = {
    "Tier 0": "Company-defining (new platform) -- Full brief + war room",
    "Tier 1": "Major feature launch -- Full brief required",
    "Tier 2": "Enhancement or update -- Light brief",
    "Tier 3": "Bug fix / minor -- No brief needed",
}

TIER_COLORS = {
    "Tier 0": ACCENT_RED,
    "Tier 1": PRIMARY_BLUE,
    "Tier 2": ACCENT_YELLOW,
    "Tier 3": ACCENT_GREEN,
}

BRIEF_TYPES = ["Campaign Brief", "Creative Brief", "GTM Brief"]

QUALITY_DIMENSIONS = [
    "Clarity",
    "Specificity",
    "Actionability",
    "Audience Focus",
    "Measurability",
]

WIZARD_STEPS = {
    1: "Setup",
    2: "Strategy",
    3: "Messaging",
    4: "Execution",
    5: "Governance",
    6: "Review & Export",
}

# ---------------------------------------------------------------------------
# Page config & custom CSS
# ---------------------------------------------------------------------------
st.set_page_config(
    page_title="Campaign Brief Generator",
    page_icon="📋",
    layout="wide",
    initial_sidebar_state="collapsed",
)

# ---------------------------------------------------------------------------
# Keep-alive: prevents WebSocket disconnect on idle browser tabs
# ---------------------------------------------------------------------------
st.components.v1.html(
    """
    <script>
    setInterval(() => {
        fetch(window.location.href, {method: 'HEAD', cache: 'no-store'});
    }, 120000);  // ping every 2 minutes
    </script>
    """,
    height=0,
)

CUSTOM_CSS = """
<style>
/* Import Inter font */
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700;800&display=swap');

/* Base */
html, body, [class*="css"] { font-family: 'Inter', sans-serif !important; }
#MainMenu {visibility: hidden;}
footer {visibility: hidden;}
header {visibility: hidden;}

/* Reduce padding */
.block-container { padding-top: 2rem; padding-bottom: 1rem; max-width: 1200px; }

/* Dark Sidebar matching Stitch #0A1628 */
section[data-testid="stSidebar"] {
    background: linear-gradient(180deg, #0A1628 0%, #132238 100%) !important;
}
section[data-testid="stSidebar"] * { color: #C8D0DA !important; }
section[data-testid="stSidebar"] h1,
section[data-testid="stSidebar"] h2,
section[data-testid="stSidebar"] h3 { color: #FFFFFF !important; }
section[data-testid="stSidebar"] hr { border-color: rgba(255,255,255,0.08) !important; }
section[data-testid="stSidebar"] .stSelectbox label,
section[data-testid="stSidebar"] .stMultiSelect label { color: #C8D0DA !important; }

/* Tab styling matching Stitch - blue underline active */
.stTabs [data-baseweb="tab-list"] {
    gap: 0;
    border-bottom: 2px solid #e0e4e8;
    background: transparent;
}
.stTabs [data-baseweb="tab"] {
    padding: 12px 24px;
    font-weight: 500;
    font-size: 0.9rem;
    color: #727785;
    border-bottom: 3px solid transparent;
    transition: all 0.2s ease;
}
.stTabs [data-baseweb="tab"]:hover {
    color: #0064D2;
    background: #E8F1FB;
}
.stTabs [aria-selected="true"] {
    color: #0064D2 !important;
    font-weight: 700 !important;
    border-bottom: 3px solid #0064D2 !important;
}

/* Step indicator badges matching Stitch */
.step-container {
    display: flex;
    justify-content: center;
    align-items: center;
    gap: 4px;
    margin-bottom: 28px;
    flex-wrap: wrap;
}
.step-badge {
    display: inline-flex;
    align-items: center;
    gap: 6px;
    padding: 8px 18px;
    border-radius: 24px;
    font-size: 0.8em;
    font-weight: 600;
    border: 2px solid #e0e4e8;
    color: #727785;
    background: #f8f9fa;
    transition: all 0.2s ease;
}
.step-badge.active {
    border-color: #004da4;
    background: #004da4;
    color: #FFFFFF;
    font-weight: 700;
    box-shadow: 0 2px 8px rgba(0,77,164,0.25);
}
.step-badge.completed {
    border-color: #507200;
    background: #f0f8e4;
    color: #507200;
    font-weight: 600;
}

/* Form inputs matching Stitch - clean rounded with blue focus */
div[data-testid="stTextInput"] input,
div[data-testid="stTextArea"] textarea {
    border-radius: 8px !important;
    border: 1.5px solid #c2c6d5 !important;
    background: #ecf5fe !important;
    transition: all 0.2s;
}
div[data-testid="stTextInput"] input:focus,
div[data-testid="stTextArea"] textarea:focus {
    border-color: #0064D2 !important;
    box-shadow: 0 0 0 3px rgba(0,100,210,0.1) !important;
}

/* Primary buttons (Next, Generate, Check Quality) - Red CTA matching Stitch secondary #ba081f */
div[data-testid="stButton"] > button[kind="primary"] {
    background: #ba081f !important;
    border: none !important;
    color: #FFFFFF !important;
    font-weight: 700 !important;
    border-radius: 8px !important;
    padding: 0.6rem 1.5rem !important;
    transition: all 0.2s ease !important;
    box-shadow: 0 4px 12px rgba(186,8,31,0.2) !important;
}
div[data-testid="stButton"] > button[kind="primary"]:hover {
    background: #930014 !important;
    box-shadow: 0 4px 20px rgba(186,8,31,0.3) !important;
    transform: translateY(-1px);
}

/* Secondary buttons (Help Me Write, Proofread, Generate, Back) - Blue outline matching Stitch */
div[data-testid="stButton"] > button[kind="secondary"],
div[data-testid="stButton"] > button:not([kind="primary"]) {
    background: transparent !important;
    border: 2px solid #004da4 !important;
    color: #004da4 !important;
    font-weight: 600 !important;
    border-radius: 8px !important;
    transition: all 0.2s ease !important;
}
div[data-testid="stButton"] > button[kind="secondary"]:hover,
div[data-testid="stButton"] > button:not([kind="primary"]):hover {
    background: rgba(0,77,164,0.05) !important;
    border-color: #004493 !important;
}

/* AI Suggestion box matching Stitch - blue left border with light blue bg */
div[data-testid="stAlert"] div[data-testid="stAlertContentInfo"] {
    background: #E8F1FB !important;
    border-left: 4px solid #0064D2 !important;
    border-radius: 0 8px 8px 0 !important;
}

/* Success box - green left border */
div[data-testid="stAlert"] div[data-testid="stAlertContentSuccess"] {
    background: #f0f8e4 !important;
    border-left: 4px solid #86B817 !important;
    border-radius: 0 8px 8px 0 !important;
}

/* Warning box - yellow */
div[data-testid="stAlert"] div[data-testid="stAlertContentWarning"] {
    background: #FFF7E0 !important;
    border-left: 4px solid #F5AF02 !important;
}

/* Error box - red */
div[data-testid="stAlert"] div[data-testid="stAlertContentError"] {
    background: #FDE8E9 !important;
    border-left: 4px solid #E53238 !important;
}

/* Quality score circle matching Stitch gradient */
.quality-score {
    display: inline-flex;
    align-items: center;
    justify-content: center;
    width: 96px;
    height: 96px;
    border-radius: 50%;
    font-size: 2.2em;
    font-weight: 800;
    color: #FFFFFF;
}
.quality-score.green  { background: linear-gradient(135deg, #86B817, #507200); box-shadow: 0 4px 12px rgba(134,184,23,0.3); }
.quality-score.yellow { background: linear-gradient(135deg, #F5AF02, #C48D02); box-shadow: 0 4px 12px rgba(245,175,2,0.3); }
.quality-score.red    { background: linear-gradient(135deg, #E53238, #930014); box-shadow: 0 4px 12px rgba(229,50,56,0.3); }

/* Tier badges */
.tier-badge {
    display: inline-block;
    padding: 3px 10px;
    border-radius: 12px;
    font-size: 0.8em;
    font-weight: 600;
    color: white;
}

/* Severity badges matching Stitch */
.severity-critical {
    display: inline-block; padding: 2px 10px; border-radius: 4px;
    background: #ffdad6; color: #93000a; font-size: 0.7em; font-weight: 700;
    text-transform: uppercase; letter-spacing: 0.5px;
}
.severity-warning {
    display: inline-block; padding: 2px 10px; border-radius: 4px;
    background: #FFF8E1; color: #F5AF02; font-size: 0.7em; font-weight: 700;
    text-transform: uppercase; letter-spacing: 0.5px;
}
.severity-suggestion {
    display: inline-block; padding: 2px 10px; border-radius: 4px;
    background: #d8e2ff; color: #004da4; font-size: 0.7em; font-weight: 700;
    text-transform: uppercase; letter-spacing: 0.5px;
}

/* Brief card styling */
.brief-card {
    border: 1px solid rgba(194,198,213,0.15);
    border-radius: 12px;
    padding: 16px;
    margin-bottom: 12px;
    background: white;
    transition: box-shadow 0.2s;
    box-shadow: 0 4px 20px rgba(0,77,164,0.06);
}
.brief-card:hover {
    box-shadow: 0 4px 12px rgba(0, 0, 0, 0.1);
}

/* Template card */
.template-card {
    border: 2px solid rgba(194,198,213,0.15);
    border-radius: 12px;
    padding: 20px;
    text-align: center;
    background: white;
    min-height: 180px;
    box-shadow: 0 4px 20px rgba(0,77,164,0.06);
}
.template-card:hover {
    border-color: #0064D2;
}

/* Pass/Fail badges */
.badge-pass {
    display: inline-block; padding: 3px 14px; border-radius: 4px;
    background: #86B817; color: #FFFFFF; font-weight: 700; font-size: 0.8em;
    text-transform: uppercase; letter-spacing: 0.5px;
}
.badge-fail {
    display: inline-block; padding: 3px 14px; border-radius: 4px;
    background: #E53238; color: #FFFFFF; font-weight: 700; font-size: 0.8em;
    text-transform: uppercase; letter-spacing: 0.5px;
}

/* Card styling matching Stitch surface-container-lowest */
.stExpander {
    border: 1px solid rgba(194,198,213,0.15) !important;
    border-radius: 12px !important;
    box-shadow: 0 4px 20px rgba(0,77,164,0.06) !important;
}

/* Selectbox styling */
div[data-testid="stSelectbox"] > div > div {
    border-radius: 8px !important;
    border-color: #c2c6d5 !important;
}

/* File uploader matching Stitch dashed border style */
div[data-testid="stFileUploader"] > div {
    border: 2px dashed rgba(194,198,213,0.3) !important;
    border-radius: 12px !important;
    background: #ffffff !important;
}
div[data-testid="stFileUploader"] > div:hover {
    border-color: #0064D2 !important;
    background: #ecf5fe !important;
}

/* Dataframe/table styling */
div[data-testid="stDataFrame"] {
    border-radius: 12px !important;
    overflow: hidden;
    box-shadow: 0 1px 3px rgba(0,0,0,0.04);
}

/* Download button styling */
div[data-testid="stDownloadButton"] > button {
    background: transparent !important;
    border: 2px solid #004da4 !important;
    color: #004da4 !important;
    font-weight: 600 !important;
    border-radius: 8px !important;
}
div[data-testid="stDownloadButton"] > button:hover {
    background: #004da4 !important;
    color: #FFFFFF !important;
}

/* Divider */
hr { border-color: rgba(194,198,213,0.15) !important; }

/* Section headers matching Stitch uppercase labels */
.section-label {
    font-size: 0.7rem;
    font-weight: 700;
    text-transform: uppercase;
    letter-spacing: 1.5px;
    color: #727785;
    margin-bottom: 4px;
}

/* Appeal bar for creative concepts */
.appeal-bar-bg {
    width: 100%; background: #e0e4e8; border-radius: 8px; height: 6px;
}
.appeal-bar-fill {
    height: 6px; border-radius: 8px;
    background: #86B817;
    transition: width 0.4s ease;
}

/* Word count badges */
.word-badge-optimal {
    display: inline-block; padding: 2px 10px; border-radius: 4px;
    background: #86B817; color: #FFFFFF; font-size: 0.75em; font-weight: 600;
}
.word-badge-long {
    display: inline-block; padding: 2px 10px; border-radius: 4px;
    background: #F5AF02; color: #FFFFFF; font-size: 0.75em; font-weight: 600;
}

/* Metric card */
.metric-card {
    background: #ffffff;
    border: 1px solid rgba(194,198,213,0.15);
    border-radius: 12px;
    padding: 20px;
    text-align: center;
    box-shadow: 0 4px 20px rgba(0,77,164,0.06);
}
.metric-value {
    font-size: 2.2em;
    font-weight: 800;
    color: #141d23;
}
.metric-label {
    font-size: 0.7rem;
    font-weight: 700;
    color: #727785;
    margin-top: 4px;
    text-transform: uppercase;
    letter-spacing: 0.5px;
}
</style>
"""

st.markdown(CUSTOM_CSS, unsafe_allow_html=True)


# ---------------------------------------------------------------------------
# Session state initialization
# ---------------------------------------------------------------------------
def init_session_state() -> None:
    """Initialize all session state variables."""
    if "wizard_step" not in st.session_state:
        st.session_state.wizard_step = 1
    if "current_brief" not in st.session_state:
        st.session_state.current_brief = {}
    if "briefs" not in st.session_state:
        st.session_state.briefs = _load_sample_briefs()
    if "demo_mode" not in st.session_state:
        api_key = os.getenv("ANTHROPIC_API_KEY", "")
        st.session_state.demo_mode = not bool(api_key)
    if "generated_names" not in st.session_state:
        st.session_state.generated_names = []
    if "ai_insight" not in st.session_state:
        st.session_state.ai_insight = None
    if "ai_positioning" not in st.session_state:
        st.session_state.ai_positioning = None
    if "ai_messages" not in st.session_state:
        st.session_state.ai_messages = None
    if "ai_smp" not in st.session_state:
        st.session_state.ai_smp = None
    if "ai_channels" not in st.session_state:
        st.session_state.ai_channels = None
    if "ai_deliverables" not in st.session_state:
        st.session_state.ai_deliverables = None
    if "ai_timeline" not in st.session_state:
        st.session_state.ai_timeline = None
    if "generated_kpis" not in st.session_state:
        st.session_state.generated_kpis = None
    if "generated_raci" not in st.session_state:
        st.session_state.generated_raci = None
    if "generated_creative_brief" not in st.session_state:
        st.session_state.generated_creative_brief = None
    if "quality_result" not in st.session_state:
        st.session_state.quality_result = None
    if "creative_concepts" not in st.session_state:
        st.session_state.creative_concepts = None
    if "audience_profile" not in st.session_state:
        st.session_state.audience_profile = None
    if "smart_objective" not in st.session_state:
        st.session_state.smart_objective = None
    # extracted_insight replaced by ai_insight above
    if "budget_breakdown" not in st.session_state:
        st.session_state.budget_breakdown = None
    if "email_sequence" not in st.session_state:
        st.session_state.email_sequence = None
    if "ai_background" not in st.session_state:
        st.session_state.ai_background = None
    if "ai_generated_objective" not in st.session_state:
        st.session_state.ai_generated_objective = None
    if "ai_generated_audience" not in st.session_state:
        st.session_state.ai_generated_audience = None


# ---------------------------------------------------------------------------
# Demo / sample data
# ---------------------------------------------------------------------------
def _load_sample_briefs() -> list[dict]:
    """Return 3 sample briefs for demo mode."""
    base = datetime(2026, 3, 30)
    return [
        {
            "id": str(uuid.uuid4()),
            "campaign_name": "AI Listing Magic",
            "brief_type": "Campaign Brief",
            "launch_tier": "Tier 1",
            "created": (base - timedelta(days=3)).isoformat(),
            "quality_score": 8.5,
            "background": (
                "Amazon recently launched an AI-powered listing generator, reducing seller "
                "onboarding friction. eBay needs to showcase its own Magical Listing feature "
                "to retain and attract sellers."
            ),
            "objective": (
                "Increase Magical Listing adoption by 40% among professional sellers "
                "within 90 days of campaign launch."
            ),
            "target_audience": (
                "Professional sellers (500+ listings) aged 25-55 who currently create "
                "listings manually. They value efficiency and listing quality."
            ),
            "key_insight": (
                "Professional sellers spend 15+ hours per week on listings -- they want "
                "more time selling, not typing."
            ),
            "positioning_short": (
                "eBay Magical Listing turns a single photo into a complete, optimized "
                "listing in seconds."
            ),
            "positioning_detailed": (
                "For professional sellers who are tired of spending hours on manual listings, "
                "eBay Magical Listing is an AI-powered tool that transforms a single product "
                "photo into a fully optimized listing with title, description, item specifics, "
                "and pricing -- saving up to 80% of listing time while improving listing "
                "quality and conversion rates compared to manual creation."
            ),
            "key_messages": [
                "1. List in seconds, not hours -- snap a photo and let AI do the rest.",
                "2. AI-optimized titles and descriptions that buyers actually search for.",
                "3. Over 10 million sellers already use Magical Listing -- join them.",
                "4. Focus on what you do best: sourcing and selling.",
            ],
            "smp": "One photo. One tap. One perfect listing.",
            "smp_pass": True,
            "channel_plan": [
                {"channel": "Email", "tactic": "Segmented drip campaign to pro sellers", "rationale": "High engagement with seller comms", "budget_pct": 25},
                {"channel": "In-App Notifications", "tactic": "Contextual prompts during listing flow", "rationale": "Catch sellers at point of need", "budget_pct": 15},
                {"channel": "Seller Hub Banners", "tactic": "Feature spotlight banner with demo video", "rationale": "Primary seller workspace", "budget_pct": 10},
                {"channel": "Social Media", "tactic": "Before/after listing demos on Instagram & TikTok", "rationale": "Visual impact, shareable content", "budget_pct": 20},
                {"channel": "YouTube", "tactic": "Tutorial series: Magical Listing mastery", "rationale": "Educational content for adoption", "budget_pct": 15},
                {"channel": "Paid Search", "tactic": "Target 'sell on eBay' and competitor queries", "rationale": "Capture high-intent sellers", "budget_pct": 15},
            ],
            "deliverables": [
                "5x email templates (welcome, tips, success story, re-engagement, upgrade)",
                "3x social video clips (15s, 30s, 60s)",
                "1x YouTube tutorial (5 min)",
                "2x Seller Hub banner designs",
                "1x landing page with demo",
                "10x in-app notification variants",
            ],
            "timeline": [
                {"phase": "Awareness", "duration": "Weeks 1-2", "activities": "Teaser emails, social posts, blog announcement"},
                {"phase": "Launch", "duration": "Weeks 3-4", "activities": "Full campaign launch, in-app prompts, YouTube video, paid search"},
                {"phase": "Sustain", "duration": "Weeks 5-8", "activities": "Success stories, re-engagement emails, optimization"},
                {"phase": "Optimize", "duration": "Weeks 9-12", "activities": "A/B test results, scale winners, reduce underperformers"},
            ],
            "kpis": [
                {"metric": "Magical Listing Adoption Rate", "target": "+40% from baseline", "measurement": "Product analytics"},
                {"metric": "Listing Creation Time", "target": "-60% reduction", "measurement": "Session analytics"},
                {"metric": "Email Open Rate", "target": ">25%", "measurement": "Email platform"},
                {"metric": "Video View Rate", "target": ">50% completion", "measurement": "YouTube / Social analytics"},
                {"metric": "New Seller Signups", "target": "+15% attributed", "measurement": "Attribution model"},
            ],
            "raci": [
                {"task": "Campaign Strategy", "responsible": "PMM Lead", "accountable": "VP Marketing", "consulted": "Product, Data", "informed": "Executives"},
                {"task": "Creative Assets", "responsible": "Creative Team", "accountable": "PMM Lead", "consulted": "Brand", "informed": "PMM"},
                {"task": "Email Execution", "responsible": "Email Marketing", "accountable": "PMM Lead", "consulted": "CRM", "informed": "Analytics"},
                {"task": "Social Campaign", "responsible": "Social Team", "accountable": "PMM Lead", "consulted": "Creative", "informed": "PR"},
                {"task": "Performance Analysis", "responsible": "Analytics", "accountable": "PMM Lead", "consulted": "Data Science", "informed": "Executives"},
            ],
            "budget": "$150,000",
        },
        {
            "id": str(uuid.uuid4()),
            "campaign_name": "Smart Ship, Happy Sellers",
            "brief_type": "GTM Brief",
            "launch_tier": "Tier 2",
            "created": (base - timedelta(days=10)).isoformat(),
            "quality_score": 7.2,
            "background": "eBay is launching improved shipping label integration with discounted USPS and UPS rates for high-volume sellers.",
            "objective": "Drive 25% increase in prepaid label usage among top 10,000 sellers within 60 days.",
            "target_audience": "High-volume sellers (1000+ monthly transactions) who currently use third-party shipping solutions.",
            "key_insight": "Sellers lose 3-5% margin to shipping inefficiencies and manual label creation.",
            "positioning_short": "eBay Shipping: The fastest, cheapest way to ship -- built right into your workflow.",
            "positioning_detailed": (
                "For high-volume sellers who waste time and money on fragmented shipping solutions, "
                "eBay's integrated shipping platform offers the lowest rates, one-click label creation, "
                "and automatic tracking -- reducing shipping costs by up to 30% and saving hours per week."
            ),
            "key_messages": [
                "1. Save up to 30% on shipping with eBay's exclusive carrier rates.",
                "2. One-click labels directly from your Seller Hub -- no copy-pasting.",
                "3. Automatic tracking updates keep buyers happy and reduce inquiries.",
            ],
            "smp": "Ship smarter, not harder.",
            "smp_pass": True,
            "channel_plan": [
                {"channel": "Email", "tactic": "Targeted campaign to top sellers", "rationale": "Direct reach to target segment", "budget_pct": 35},
                {"channel": "Seller Hub", "tactic": "Interactive shipping calculator widget", "rationale": "Show savings in context", "budget_pct": 25},
                {"channel": "Webinar", "tactic": "Live demo with Q&A", "rationale": "Build trust and answer objections", "budget_pct": 20},
                {"channel": "Community", "tactic": "Seller forum posts and success stories", "rationale": "Peer validation", "budget_pct": 20},
            ],
            "deliverables": [
                "3x email templates",
                "1x shipping calculator widget",
                "1x webinar deck and recording",
                "5x community posts",
            ],
            "timeline": [
                {"phase": "Awareness", "duration": "Week 1", "activities": "Email announcement, forum posts"},
                {"phase": "Launch", "duration": "Weeks 2-3", "activities": "Webinar, calculator launch, Seller Hub banners"},
                {"phase": "Sustain", "duration": "Weeks 4-8", "activities": "Success stories, optimization"},
            ],
            "kpis": [
                {"metric": "Prepaid Label Usage", "target": "+25%", "measurement": "Product analytics"},
                {"metric": "Shipping Cost Reduction", "target": "-20% avg", "measurement": "Transaction data"},
                {"metric": "Webinar Attendance", "target": "500+ attendees", "measurement": "Webinar platform"},
            ],
            "raci": [
                {"task": "GTM Strategy", "responsible": "PMM Lead", "accountable": "Dir. Marketing", "consulted": "Product", "informed": "Shipping Ops"},
                {"task": "Webinar", "responsible": "Content Team", "accountable": "PMM Lead", "consulted": "Product", "informed": "Sales"},
            ],
            "budget": "$75,000",
        },
        {
            "id": str(uuid.uuid4()),
            "campaign_name": "Collect with Confidence",
            "brief_type": "Creative Brief",
            "launch_tier": "Tier 1",
            "created": (base - timedelta(days=20)).isoformat(),
            "quality_score": 6.1,
            "background": "eBay's Authenticity Guarantee program has expanded to cover more collectible categories. Need to drive awareness among collectors.",
            "objective": "Increase Authenticity Guarantee awareness to 60% among active collectors and drive 20% more purchases in covered categories.",
            "target_audience": "Collectors aged 30-60 who spend $500+ annually on collectibles (trading cards, coins, sneakers, watches).",
            "key_insight": "Collectors fear fakes more than high prices -- trust is the #1 purchase barrier.",
            "positioning_short": "eBay Authenticity Guarantee: Every item verified, every purchase protected.",
            "positioning_detailed": (
                "For passionate collectors who fear counterfeit goods, eBay's Authenticity Guarantee "
                "provides expert verification on every eligible item -- giving buyers absolute confidence "
                "that what they collect is real, protected, and worth every penny."
            ),
            "key_messages": [
                "1. Every item inspected by independent experts before it reaches you.",
                "2. If it's not authentic, you get your money back -- guaranteed.",
                "3. eBay is the only marketplace offering free authentication at this scale.",
            ],
            "smp": "Real items. Real experts. Real confidence.",
            "smp_pass": True,
            "channel_plan": [
                {"channel": "Social Media", "tactic": "Unboxing videos with authentication reveal", "rationale": "Visual trust-building", "budget_pct": 30},
                {"channel": "Influencer", "tactic": "Partner with top collectors on YouTube", "rationale": "Credibility through community leaders", "budget_pct": 25},
                {"channel": "Email", "tactic": "Category-specific trust campaigns", "rationale": "Personalized relevance", "budget_pct": 20},
                {"channel": "PR", "tactic": "Media stories on authentication process", "rationale": "Third-party validation", "budget_pct": 25},
            ],
            "deliverables": [
                "4x unboxing video scripts",
                "3x influencer partnership briefs",
                "6x email templates by category",
                "2x press releases",
                "1x behind-the-scenes authentication video",
            ],
            "timeline": [
                {"phase": "Awareness", "duration": "Weeks 1-3", "activities": "Influencer content, PR push, teaser social"},
                {"phase": "Launch", "duration": "Weeks 4-5", "activities": "Full campaign launch across channels"},
                {"phase": "Sustain", "duration": "Weeks 6-10", "activities": "Category deep-dives, retargeting"},
            ],
            "kpis": [
                {"metric": "AG Awareness", "target": "60% among collectors", "measurement": "Survey"},
                {"metric": "AG Category GMV", "target": "+20%", "measurement": "Transaction data"},
                {"metric": "Video Views", "target": "1M+ total", "measurement": "Social analytics"},
            ],
            "raci": [
                {"task": "Creative Direction", "responsible": "Creative Lead", "accountable": "Brand Dir.", "consulted": "PMM", "informed": "Product"},
                {"task": "Influencer Mgmt", "responsible": "Social Team", "accountable": "PMM Lead", "consulted": "Legal", "informed": "PR"},
            ],
            "budget": "$200,000",
        },
    ]


def _get_demo_quality_result() -> dict:
    """Return a sample quality analysis for demo mode."""
    return {
        "overall_score": 7.8,
        "letter_grade": "B+",
        "dimensions": {
            "Clarity": 8.5,
            "Specificity": 7.0,
            "Actionability": 8.0,
            "Audience Focus": 7.5,
            "Measurability": 8.0,
        },
        "mistakes": [
            {"issue": "Multiple objectives detected -- brief has both adoption and awareness goals", "severity": "warning", "section": "Objective"},
            {"issue": "Target audience could be more specific -- add psychographic details", "severity": "suggestion", "section": "Target Audience"},
            {"issue": "No competitive context mentioned in background", "severity": "warning", "section": "Background"},
        ],
        "missing_sections": ["Budget Breakdown", "Risk Assessment"],
        "strengths": [
            "Clear SMART objective with measurable target",
            "Well-defined channel plan with rationale",
            "Strong SMP -- single-minded and memorable",
            "Phased timeline with clear activities",
        ],
        "improvements": [
            "1. Narrow to a single primary objective. Move the secondary goal to a stretch target.",
            "2. Add psychographic details to the target audience (values, behaviors, media habits).",
            "3. Include competitive context: what are Amazon/Shopify doing in this space?",
            "4. Add a risk assessment section with mitigation strategies.",
            "5. Include a budget breakdown even if approximate.",
        ],
        "smp_check": {
            "smp": "One photo. One tap. One perfect listing.",
            "pass": True,
            "reason": "SMP is genuinely single-minded, focusing on simplicity. No 'and/also' detected.",
        },
    }


def _get_demo_concepts() -> list[dict]:
    """Return 3 sample creative concepts for demo mode."""
    return [
        {
            "name": "The Listing Revolution",
            "theme": "Empowerment through technology",
            "headline": "Stop Typing. Start Selling.",
            "tagline": "One photo changes everything.",
            "tone": "Bold, confident, empowering",
            "appeal_score": 8.5,
            "sample_copy": (
                "You didn't become a seller to spend hours writing descriptions. "
                "With eBay Magical Listing, snap a photo and watch AI create a "
                "perfectly optimized listing in seconds. It's not just faster -- "
                "it's better than manual."
            ),
            "visual_direction": (
                "Split-screen concept: left shows cluttered desk with papers and "
                "a frustrated seller; right shows clean, modern workspace with a "
                "phone and a perfect listing. Bright, energetic color palette."
            ),
        },
        {
            "name": "Time is Money",
            "theme": "Efficiency and ROI",
            "headline": "15 Hours a Week. Back in Your Pocket.",
            "tagline": "AI lists it. You profit from it.",
            "tone": "Professional, data-driven, persuasive",
            "appeal_score": 7.8,
            "sample_copy": (
                "Professional sellers spend 15+ hours weekly on listings. "
                "eBay Magical Listing cuts that to minutes. That's 780 hours "
                "a year you could spend sourcing, strategizing, or just living. "
                "The math is simple."
            ),
            "visual_direction": (
                "Infographic-style visuals with a ticking clock motif. "
                "Show the 'time saved' accumulating like a savings account. "
                "Clean typography, data visualizations, professional blue palette."
            ),
        },
        {
            "name": "Photo Finish",
            "theme": "Simplicity and magic",
            "headline": "Your Phone Already Has the Perfect Listing.",
            "tagline": "Snap. List. Sell.",
            "tone": "Playful, approachable, magical",
            "appeal_score": 9.0,
            "sample_copy": (
                "That photo you just took? It's already a complete listing -- "
                "title, description, price, and all. eBay's AI sees what you see, "
                "then writes what buyers want to read. It feels like magic "
                "because it kind of is."
            ),
            "visual_direction": (
                "Whimsical, almost magical transformation sequence. Phone camera "
                "flash leads to sparkle effects as a listing materializes. "
                "Warm, inviting tones with pops of eBay blue. Feels like an "
                "Apple product reveal meets a street magic show."
            ),
        },
    ]


def _get_demo_email_sequence() -> list[dict]:
    """Return a 5-email drip campaign for demo mode."""
    return [
        {
            "email_num": 1,
            "subject": "Your listings are about to get a whole lot easier",
            "day": "Day 0",
            "purpose": "Welcome & Introduction",
            "preview": (
                "Hi [Name],\n\n"
                "We noticed you create a lot of listings manually. What if we told you "
                "there's a faster way?\n\n"
                "Introducing Magical Listing -- snap a photo, and AI does the rest. "
                "Title, description, item specifics, pricing -- all optimized for search.\n\n"
                "Try it now on your next listing."
            ),
        },
        {
            "email_num": 2,
            "subject": "3 tips to get the most from Magical Listing",
            "day": "Day 3",
            "purpose": "Education & Tips",
            "preview": (
                "Great lighting, multiple angles, and a clean background -- that's all "
                "it takes to get a perfect AI-generated listing. Here are 3 pro tips "
                "from our top sellers..."
            ),
        },
        {
            "email_num": 3,
            "subject": "How Sarah went from 50 to 200 listings per week",
            "day": "Day 7",
            "purpose": "Social Proof",
            "preview": (
                "Sarah runs a vintage clothing shop on eBay. She was spending 20 hours "
                "a week on listings. Now? Four hours. Here's her story..."
            ),
        },
        {
            "email_num": 4,
            "subject": "You haven't tried Magical Listing yet -- here's why you should",
            "day": "Day 14",
            "purpose": "Re-engagement",
            "preview": (
                "We get it -- change is hard. But 10 million sellers have already switched "
                "to AI-powered listings. Here's what they found..."
            ),
        },
        {
            "email_num": 5,
            "subject": "Exclusive: Bulk Listing AI now available for pro sellers",
            "day": "Day 21",
            "purpose": "Upgrade & Upsell",
            "preview": (
                "Love Magical Listing? Meet Bulk Listing AI -- upload 50 photos at once "
                "and get 50 optimized listings in minutes. Available now for pro sellers."
            ),
        },
    ]


SAMPLE_BRIEF_TEXT = """Campaign Brief: AI-Powered Listing Generator

Background: Amazon recently launched an AI listing tool. We need to promote eBay's Magical Listing feature.

Objective: Increase adoption by 40% among professional sellers in Q2.

Target Audience: Professional sellers on eBay.

Key Messages:
1. List faster with AI
2. Better listings mean more sales
3. Join millions of sellers using Magical Listing

Timeline: Q2 2026

Channels: Email, social media, in-app
"""

DEMO_PREFILLED_BRIEF = {
    "campaign_name": "AI-Powered Listing Generator Launch",
    "launch_tier": "Tier 1",
    "brief_type": "Campaign Brief",
    "background": (
        "Amazon recently launched an AI-powered listing generator, reducing seller "
        "onboarding friction. eBay needs to showcase its own Magical Listing feature "
        "to retain and attract professional sellers who are evaluating competitive "
        "platforms. The timing is critical as Q2 is peak seller acquisition season."
    ),
    "objective": (
        "Increase Magical Listing adoption by 40% among professional sellers "
        "(500+ active listings) within 90 days of campaign launch in Q2 2026."
    ),
    "target_audience": (
        "Professional sellers aged 25-55 with 500+ active listings who currently "
        "create listings manually. They value time efficiency, listing quality, "
        "and tools that scale with their business."
    ),
    "key_insight": (
        "Professional sellers spend 15+ hours per week on listing creation -- "
        "they want more time selling, not typing descriptions."
    ),
}


# ---------------------------------------------------------------------------
# Helper functions
# ---------------------------------------------------------------------------


def _format_audience_profile(profile) -> str:
    """Format an audience profile dict into readable markdown."""
    if isinstance(profile, str):
        return profile
    if not isinstance(profile, dict):
        return str(profile)

    lines = []
    if profile.get("demographics"):
        lines.append(f"**Demographics:** {profile['demographics']}\n")
    if profile.get("psychographics"):
        lines.append(f"**Psychographics:** {profile['psychographics']}\n")
    if profile.get("pain_points"):
        lines.append("**Pain Points:**")
        for p in profile["pain_points"]:
            lines.append(f"- {p}")
        lines.append("")
    if profile.get("motivations"):
        lines.append("**Motivations:**")
        for m in profile["motivations"]:
            lines.append(f"- {m}")
        lines.append("")
    if profile.get("media_habits"):
        lines.append("**Media Habits:**")
        for h in profile["media_habits"]:
            lines.append(f"- {h}")
        lines.append("")
    if profile.get("objections"):
        lines.append("**Likely Objections:**")
        for o in profile["objections"]:
            lines.append(f"- {o}")
        lines.append("")
    return "\n".join(lines)


def _format_creative_brief(cb) -> str:
    """Format a creative brief dict into readable markdown."""
    if isinstance(cb, str):
        return cb
    if not isinstance(cb, dict):
        return str(cb)

    lines = []
    lines.append(f"# Creative Brief: {cb.get('project_name', 'Untitled')}\n")
    if cb.get("objective"):
        lines.append(f"## Objective\n{cb['objective']}\n")
    if cb.get("target_audience"):
        lines.append(f"## Target Audience\n{cb['target_audience']}\n")
    if cb.get("single_minded_proposition"):
        lines.append(f"## Single-Minded Proposition\n*{cb['single_minded_proposition']}*\n")
    if cb.get("tone_and_manner"):
        lines.append(f"## Tone & Manner\n{cb['tone_and_manner']}\n")
    if cb.get("mandatories"):
        lines.append("## Mandatories")
        for m in cb["mandatories"]:
            lines.append(f"- {m}")
        lines.append("")
    if cb.get("deliverables"):
        lines.append("## Deliverables")
        for d in cb["deliverables"]:
            if isinstance(d, dict):
                lines.append(f"- **{d.get('asset', '')}**: {d.get('specs', d.get('spec', ''))} ({d.get('word_count', '')})")
            else:
                lines.append(f"- {d}")
        lines.append("")
    if cb.get("inspiration_references"):
        lines.append("## Inspiration & References")
        for r in cb["inspiration_references"]:
            lines.append(f"- {r}")
        lines.append("")
    if cb.get("do_nots"):
        lines.append("## Do Nots")
        for d in cb["do_nots"]:
            lines.append(f"- {d}")
        lines.append("")
    return "\n".join(lines)


def _get_generator() -> Optional[Any]:
    """Get BriefGenerator instance, or None if in demo mode."""
    if st.session_state.demo_mode:
        return None
    try:
        return BriefGenerator()
    except Exception as e:
        logger.error("Failed to initialize BriefGenerator: %s", e)
        return None


def _get_quality_checker() -> Optional[Any]:
    """Get QualityChecker instance, or None if in demo mode."""
    if st.session_state.demo_mode:
        return None
    try:
        return QualityChecker()
    except Exception as e:
        logger.error("Failed to initialize QualityChecker: %s", e)
        return None


def _get_creative_engine() -> Optional[Any]:
    """Get CreativeEngine instance, or None if in demo mode."""
    if st.session_state.demo_mode:
        return None
    try:
        return CreativeEngine()
    except Exception as e:
        logger.error("Failed to initialize CreativeEngine: %s", e)
        return None


def render_step_indicators(current_step: int) -> None:
    """Render the wizard step indicator bar."""
    html_parts = ['<div class="step-container">']
    for step_num, step_name in WIZARD_STEPS.items():
        if step_num < current_step:
            cls = "step-badge completed"
            icon = "&#10003;"
        elif step_num == current_step:
            cls = "step-badge active"
            icon = str(step_num)
        else:
            cls = "step-badge"
            icon = str(step_num)
        html_parts.append(
            f'<span class="{cls}">{icon} {step_name}</span>'
        )
    html_parts.append("</div>")
    st.markdown("".join(html_parts), unsafe_allow_html=True)


def render_tier_badge(tier: str) -> str:
    """Return HTML for a tier badge."""
    color = TIER_COLORS.get(tier, "#888")
    return f'<span class="tier-badge" style="background:{color}">{tier}</span>'


def render_quality_score_badge(score: float) -> str:
    """Return HTML for a quality score circle badge."""
    if score > 7:
        cls = "green"
    elif score >= 5:
        cls = "yellow"
    else:
        cls = "red"
    return f'<div class="quality-score {cls}">{score:.1f}</div>'


def render_severity_badge(severity: str) -> str:
    """Return HTML for a severity badge."""
    cls = f"severity-{severity}"
    return f'<span class="{cls}">{severity.upper()}</span>'


def brief_to_markdown(brief: dict) -> str:
    """Convert a brief dict to a markdown string for preview/export."""
    lines = []
    lines.append(f"# Campaign Brief: {brief.get('campaign_name', 'Untitled')}")
    lines.append("")
    lines.append(f"**Type:** {brief.get('brief_type', 'N/A')}  ")
    lines.append(f"**Launch Tier:** {brief.get('launch_tier', 'N/A')}  ")
    lines.append(f"**Date:** {brief.get('created', datetime.now().isoformat())[:10]}")
    lines.append("")

    lines.append("---")
    lines.append("## Background & Context")
    lines.append(brief.get("background", "Not specified"))
    lines.append("")

    lines.append("## Objective")
    lines.append(brief.get("objective", "Not specified"))
    lines.append("")

    lines.append("## Target Audience")
    lines.append(brief.get("target_audience", "Not specified"))
    lines.append("")

    lines.append("## Key Insight")
    lines.append(brief.get("key_insight", "Not specified"))
    lines.append("")

    if brief.get("positioning_short"):
        lines.append("## Positioning Statement")
        lines.append(f"**Short (25 words):** {brief['positioning_short']}")
        lines.append("")
        if brief.get("positioning_detailed"):
            lines.append(f"**Detailed:** {brief['positioning_detailed']}")
            lines.append("")

    if brief.get("key_messages"):
        lines.append("## Key Messages")
        for msg in brief["key_messages"]:
            lines.append(f"- {msg}")
        lines.append("")

    if brief.get("smp"):
        lines.append("## Single-Minded Proposition (SMP)")
        lines.append(f"**{brief['smp']}**")
        if brief.get("smp_pass") is not None:
            status = "PASS" if brief["smp_pass"] else "FAIL"
            lines.append(f"  *Quality Check: {status}*")
        lines.append("")

    if brief.get("channel_plan"):
        lines.append("## Channel Plan")
        lines.append("| Channel | Tactic | Rationale | Budget % |")
        lines.append("|---------|--------|-----------|----------|")
        for ch in brief["channel_plan"]:
            lines.append(
                f"| {ch.get('channel','')} | {ch.get('tactic','')} | "
                f"{ch.get('rationale','')} | {ch.get('budget_pct','')}% |"
            )
        lines.append("")

    if brief.get("deliverables"):
        lines.append("## Content Deliverables")
        lines.append("| Asset | Spec | Owner |")
        lines.append("|-------|------|-------|")
        for d in brief["deliverables"]:
            if isinstance(d, dict):
                lines.append(
                    f"| {d.get('asset', '')} | {d.get('spec', '')} | {d.get('owner', '')} |"
                )
            else:
                lines.append(f"| {d} | | |")
        lines.append("")

    if brief.get("timeline"):
        lines.append("## Timeline")
        lines.append("| Phase | Duration | Activities |")
        lines.append("|-------|----------|------------|")
        for t in brief["timeline"]:
            lines.append(
                f"| {t.get('phase','')} | {t.get('duration','')} | {t.get('activities','')} |"
            )
        lines.append("")

    if brief.get("kpis"):
        lines.append("## Success Metrics / KPIs")
        lines.append("| Metric | Target | Measurement |")
        lines.append("|--------|--------|-------------|")
        for k in brief["kpis"]:
            lines.append(
                f"| {k.get('metric','')} | {k.get('target','')} | {k.get('measurement','')} |"
            )
        lines.append("")

    if brief.get("raci"):
        lines.append("## RACI Matrix")
        lines.append("| Task | Responsible | Accountable | Consulted | Informed |")
        lines.append("|------|-------------|-------------|-----------|----------|")
        for r in brief["raci"]:
            lines.append(
                f"| {r.get('task','')} | {r.get('responsible','')} | "
                f"{r.get('accountable','')} | {r.get('consulted','')} | "
                f"{r.get('informed','')} |"
            )
        lines.append("")

    if brief.get("budget"):
        lines.append("## Budget")
        lines.append(f"**Total:** {brief['budget']}")
        lines.append("")

    return "\n".join(lines)


# ---------------------------------------------------------------------------
# Tab 1: AI Brief Builder
# ---------------------------------------------------------------------------
def render_brief_builder() -> None:
    """Render the AI Brief Builder wizard tab."""
    step = st.session_state.wizard_step
    render_step_indicators(step)

    if st.session_state.demo_mode:
        # Pre-fill with demo data if brief is empty
        if not st.session_state.current_brief:
            st.session_state.current_brief = DEMO_PREFILLED_BRIEF.copy()

    brief = st.session_state.current_brief
    generator = _get_generator()

    # --- Step 1: Setup ---
    if step == 1:
        st.subheader("Step 1: Campaign Setup")
        col1, col2 = st.columns([3, 1])
        with col1:
            campaign_name = st.text_input(
                "Campaign Name",
                value=brief.get("campaign_name", ""),
                placeholder="e.g., AI Listing Magic Launch",
            )
        with col2:
            st.write("")
            st.write("")
            if st.button("Suggest Names", use_container_width=True):
                if generator and not st.session_state.demo_mode:
                    with st.spinner("Generating name suggestions..."):
                        try:
                            names = generator.generate_campaign_names(
                                campaign_name or brief.get("background", "AI feature launch"),
                                brief.get("target_audience", "") or "e-commerce sellers",
                            )
                            st.session_state.generated_names = names
                        except Exception as e:
                            st.error(f"Error: {e}")
                else:
                    st.session_state.generated_names = [
                        "AI Listing Magic",
                        "List Smarter Launch",
                        "The Listing Revolution",
                        "Snap, List, Sell",
                        "Magical Listing Blitz",
                    ]

        if st.session_state.generated_names:
            col_names_info, col_names_close = st.columns([11, 1])
            with col_names_info:
                st.info("**Suggested names:** " + " | ".join(st.session_state.generated_names))
            with col_names_close:
                st.write("")
                if st.button("Close X", key="close_names"):
                    st.session_state.generated_names = []
                    st.rerun()

        tier_options = [""] + list(TIER_DESCRIPTIONS.keys())
        current_tier = brief.get("launch_tier", "")
        tier_index = tier_options.index(current_tier) if current_tier in tier_options else 0
        launch_tier = st.selectbox(
            "Select Launch Tier",
            options=tier_options,
            index=tier_index,
            format_func=lambda t: "Select a launch tier..." if t == "" else f"{t} -- {TIER_DESCRIPTIONS[t]}",
            help=(
                "**Tier 0 — Company-defining:** New platform or company-wide initiative. "
                "Requires a full 14-section brief plus a dedicated war room.\n\n"
                "**Tier 1 — Major feature launch:** Significant new capability. "
                "Requires a full campaign brief with cross-functional alignment.\n\n"
                "**Tier 2 — Enhancement or update:** Improvement to an existing feature. "
                "A lighter brief covering key sections is sufficient.\n\n"
                "**Tier 3 — Bug fix / minor:** Small change with minimal marketing impact. "
                "Typically no brief needed."
            ),
        )

        brief_type = st.radio(
            "Brief Type",
            BRIEF_TYPES,
            index=BRIEF_TYPES.index(brief.get("brief_type", "Campaign Brief")),
            horizontal=True,
            help=(
                "**Campaign Brief:** Strategic document covering the full 'what & why' "
                "-- objectives, audience, channels, budget, KPIs. Used to align cross-functional teams.\n\n"
                "**Creative Brief:** Execution doc for designers/agencies -- tone, visuals, "
                "mandatories, deliverable specs. Created AFTER the campaign brief.\n\n"
                "**GTM Brief:** Launch-specific doc including pricing, sales enablement, "
                "training plans, competitive positioning. For new product/feature launches."
            ),
        )

        st.write("")
        if st.button("Next  >>", type="primary", use_container_width=True):
            # Validation
            missing = []
            if not campaign_name.strip():
                missing.append("Campaign Name")
            if launch_tier == "":
                missing.append("Launch Tier")
            if missing:
                st.error(f"Please fill in all required fields: {', '.join(missing)}")
            else:
                brief["campaign_name"] = campaign_name
                brief["launch_tier"] = launch_tier
                brief["brief_type"] = brief_type
                st.session_state.current_brief = brief
                st.session_state.wizard_step = 2
                st.rerun()

    # --- Step 2: Strategy ---
    elif step == 2:
        st.subheader("Step 2: Strategy")

        background = st.text_area(
            "Background / Context",
            value=brief.get("background", ""),
            height=120,
            placeholder="Why is this campaign needed now? What market conditions or competitive moves prompted it?",
        )

        uploaded_files = st.file_uploader(
            "Upload supporting documents (optional)",
            type=["pdf", "txt", "csv", "md", "docx"],
            accept_multiple_files=True,
            key="background_docs",
            help="Upload market research, competitive analysis, product specs, or any context documents.",
        )
        doc_content = ""
        if uploaded_files:
            extracted_texts = []
            for uf in uploaded_files:
                try:
                    if uf.name.endswith(".pdf"):
                        try:
                            import PyPDF2
                            reader = PyPDF2.PdfReader(uf)
                            text = "\n".join(page.extract_text() or "" for page in reader.pages[:10])
                        except ImportError:
                            text = "[PDF support requires PyPDF2: pip install PyPDF2]"
                    elif uf.name.endswith(".docx"):
                        try:
                            import docx
                            doc = docx.Document(uf)
                            text = "\n".join(p.text for p in doc.paragraphs)
                        except ImportError:
                            text = "[DOCX support requires python-docx: pip install python-docx]"
                    else:
                        text = uf.read().decode("utf-8", errors="ignore")
                    extracted_texts.append(f"--- {uf.name} ---\n{text[:3000]}")
                except Exception as e:
                    extracted_texts.append(f"--- {uf.name} ---\n[Error reading file: {e}]")

            if extracted_texts:
                doc_content = "\n\n".join(extracted_texts)
                with st.expander(f"Extracted content from {len(uploaded_files)} document(s)", expanded=False):
                    st.text(doc_content[:2000] + ("..." if len(doc_content) > 2000 else ""))

        # Combine background with uploaded doc content for downstream use
        full_background = background
        if doc_content.strip():
            full_background = background + "\n\n--- Uploaded Documents ---\n" + doc_content if background else doc_content

        # --- AI writing assistance buttons for Background ---
        col_help_bg, col_proof_bg = st.columns(2)
        with col_help_bg:
            if st.button("Help Me Write", key="bg_help", use_container_width=True,
                         help="Expands your notes and uploaded docs into a full background paragraph"):
                if not background.strip() and not doc_content:
                    st.warning("Enter some notes in Background/Context or upload documents first, then click Help Me Write.")
                else:
                    campaign_name_val = brief.get("campaign_name", "")
                    input_context = background + ("\n\n" + doc_content if doc_content else "")
                    if generator and not st.session_state.demo_mode:
                        with st.spinner("Expanding your notes into a full background..."):
                            try:
                                result = generator.help_write_background(campaign_name_val, input_context)
                                st.session_state.ai_background = result
                            except Exception as e:
                                st.error(f"Error: {e}")
                    else:
                        st.session_state.ai_background = (
                            f"Based on your input about '{background[:50]}...', here is an expanded background:\n\n"
                            f"The e-commerce landscape is rapidly evolving as platforms compete to simplify "
                            f"the seller experience. {brief.get('campaign_name', 'This initiative')} addresses "
                            f"a critical gap in the current workflow. Competitors have recently launched "
                            f"similar capabilities, raising customer expectations. Our internal data shows "
                            f"growing demand for this functionality, making the timing critical to capture "
                            f"mindshare before competitors consolidate their position."
                        )
                    st.rerun()
        with col_proof_bg:
            if st.button("Proofread", key="bg_proof", use_container_width=True,
                         help="Refines your background text for grammar, clarity, and professional tone"):
                if not background.strip():
                    st.warning("Write some background text first, then click Proofread.")
                else:
                    if generator and not st.session_state.demo_mode:
                        with st.spinner("Proofreading..."):
                            try:
                                result = generator.proofread_text(background)
                                st.session_state.ai_background = result
                            except Exception as e:
                                st.error(f"Error: {e}")
                    else:
                        st.session_state.ai_background = (
                            background.rstrip(".") + ". "
                            "[Proofread] This text has been refined for clarity, grammar, "
                            "and professional tone. Sentences have been tightened and "
                            "passive voice converted to active voice."
                        )
                    st.rerun()

        if st.session_state.get("ai_background"):
            col_bg_info, col_bg_close = st.columns([11, 1])
            with col_bg_info:
                st.info(st.session_state.ai_background)
            with col_bg_close:
                st.write("")
                if st.button("Close X", key="close_bg"):
                    st.session_state.ai_background = None
                    st.rerun()
            if st.button("Use This", key="use_bg"):
                brief["background"] = st.session_state.ai_background
                st.session_state.ai_background = None
                st.session_state.current_brief = brief
                st.rerun()

        # --- Objective (full width) with buttons below ---
        objective = st.text_input(
            "Objective",
            value=brief.get("objective", st.session_state.smart_objective or ""),
            placeholder="What is the single most important goal?",
        )

        col_gen_obj, col_smart = st.columns(2)
        with col_smart:
            if st.button("Make it SMART", use_container_width=True, key="btn_smart"):
                bg_context = f" Context: {full_background[:300]}" if full_background else ""
                if objective.strip():
                    smart_input = f"{objective.strip()}.{bg_context}"
                else:
                    smart_input = (
                        f"Create a campaign objective for: {brief.get('campaign_name', 'feature launch')}."
                        f"{bg_context}"
                    )
                if generator and not st.session_state.demo_mode:
                    with st.spinner("Generating SMART objective..."):
                        try:
                            smart = generator.make_objective_smart(smart_input)
                            st.session_state.smart_objective = smart
                        except Exception as e:
                            st.error(f"Error: {e}")
                else:
                    st.session_state.smart_objective = (
                        "Increase Magical Listing adoption by 40% (from 25% to 35% of total listings) "
                        "among professional sellers (500+ active listings) within 90 days of "
                        "campaign launch, measured via product analytics dashboard."
                    )
                st.rerun()
        with col_gen_obj:
            if st.button("Generate Objective", use_container_width=True, key="btn_gen_obj"):
                campaign_name_val = brief.get("campaign_name", "")
                if generator and not st.session_state.demo_mode:
                    with st.spinner("Generating objective..."):
                        try:
                            obj = generator.generate_objective(campaign_name_val, full_background)
                            st.session_state.ai_generated_objective = obj
                        except Exception as e:
                            st.error(f"Error: {e}")
                else:
                    st.session_state.ai_generated_objective = (
                        f"Increase {campaign_name_val or 'feature'} adoption by 40% among professional "
                        "users (those with 50+ monthly transactions) within 90 days of campaign launch "
                        "in Q2 2026, measured through product analytics dashboard with weekly progress "
                        "tracking against the pre-campaign baseline."
                    )
                st.rerun()

        if st.session_state.smart_objective:
            col_smart_info, col_smart_close = st.columns([11, 1])
            with col_smart_info:
                st.success(f"**SMART Objective:** {st.session_state.smart_objective}")
            with col_smart_close:
                st.write("")
                if st.button("Close X", key="close_smart"):
                    st.session_state.smart_objective = None
                    st.rerun()
            if st.button("Use This", key="use_smart"):
                brief["objective"] = st.session_state.smart_objective
                st.session_state.current_brief = brief
                st.session_state.smart_objective = None
                st.rerun()

        if st.session_state.get("ai_generated_objective"):
            col_obj_info, col_obj_close = st.columns([11, 1])
            with col_obj_info:
                st.info(f"**Generated Objective:** {st.session_state.ai_generated_objective}")
            with col_obj_close:
                st.write("")
                if st.button("Close X", key="close_obj"):
                    st.session_state.ai_generated_objective = None
                    st.rerun()
            if st.button("Use This Objective", key="use_obj"):
                brief["objective"] = st.session_state.ai_generated_objective
                st.session_state.smart_objective = st.session_state.ai_generated_objective
                st.session_state.ai_generated_objective = None
                st.session_state.current_brief = brief
                st.rerun()

        # --- Target Audience (full width) with buttons below ---
        target_audience = st.text_area(
            "Target Audience",
            value=brief.get("target_audience", ""),
            height=100,
            placeholder="Who are we talking to? Be specific about demographics, behaviors, and needs.",
        )

        col_profile, col_help_aud = st.columns(2)
        with col_profile:
            if st.button("Generate Profile", use_container_width=True, key="btn_gen_profile"):
                audience_input = target_audience.strip() if target_audience.strip() else (
                    f"Target audience for: {brief.get('campaign_name', 'campaign')}. "
                    f"Context: {full_background[:200] if full_background else 'e-commerce feature launch'}. "
                    f"Objective: {brief.get('objective', objective or 'drive adoption')}"
                )
                if generator and not st.session_state.demo_mode:
                    with st.spinner("Generating audience profile..."):
                        try:
                            profile = generator.generate_audience_profile(audience_input)
                            st.session_state.audience_profile = _format_audience_profile(profile)
                        except Exception as e:
                            st.error(f"Error: {e}")
                else:
                    st.session_state.audience_profile = (
                        "**Primary Persona: The Efficiency-Driven Pro Seller**\n\n"
                        "- **Demographics:** Ages 25-55, running small-to-medium e-commerce businesses\n"
                        "- **Behavior:** Lists 50-200 items/week, spends 15+ hours on listing creation\n"
                        "- **Pain Points:** Manual listing is tedious, inconsistent quality, poor SEO\n"
                        "- **Motivation:** Save time, improve listing quality, increase sales velocity\n"
                        "- **Media Habits:** Seller Hub daily, YouTube tutorials, seller community forums\n"
                        "- **Decision Drivers:** ROI proof, ease of use, peer recommendations"
                    )
                st.rerun()
        with col_help_aud:
            if st.button("Help Me Write", key="aud_help", use_container_width=True):
                campaign_name_val = brief.get("campaign_name", "")
                obj_val = st.session_state.smart_objective or objective or brief.get("objective", "")
                if generator and not st.session_state.demo_mode:
                    with st.spinner("Generating audience description..."):
                        try:
                            aud = generator.help_write_audience(campaign_name_val, full_background, obj_val)
                            st.session_state.ai_generated_audience = aud
                        except Exception as e:
                            st.error(f"Error: {e}")
                else:
                    st.session_state.ai_generated_audience = (
                        f"Primary audience for {campaign_name_val or 'this campaign'}: Professional users "
                        "aged 25-55 who are active on the platform (50+ monthly transactions). They are "
                        "efficiency-driven, tech-savvy early adopters who value tools that save time and "
                        "scale their operations. Key pain points include manual workflows, inconsistent "
                        "quality, and difficulty scaling beyond current volume."
                    )
                st.rerun()

        if st.session_state.audience_profile:
            col_prof_info, col_prof_close = st.columns([11, 1])
            with col_prof_info:
                with st.expander("Generated Audience Profile", expanded=True):
                    st.markdown(st.session_state.audience_profile)
            with col_prof_close:
                st.write("")
                if st.button("Close X", key="close_profile"):
                    st.session_state.audience_profile = None
                    st.rerun()
            if st.button("Use This", key="use_profile"):
                brief["target_audience"] = st.session_state.audience_profile
                st.session_state.current_brief = brief
                st.session_state.audience_profile = None
                st.rerun()

        if st.session_state.get("ai_generated_audience"):
            col_aud_info, col_aud_close = st.columns([11, 1])
            with col_aud_info:
                st.info(st.session_state.ai_generated_audience)
            with col_aud_close:
                st.write("")
                if st.button("Close X", key="close_aud"):
                    st.session_state.ai_generated_audience = None
                    st.rerun()
            if st.button("Use This Audience", key="use_aud"):
                brief["target_audience"] = st.session_state.ai_generated_audience
                st.session_state.audience_profile = st.session_state.ai_generated_audience
                st.session_state.ai_generated_audience = None
                st.session_state.current_brief = brief
                st.rerun()

        st.write("")
        col_back, col_next = st.columns(2)
        with col_back:
            if st.button("<<  Back", use_container_width=True):
                st.session_state.wizard_step = 1
                st.rerun()
        with col_next:
            if st.button("Next  >>", type="primary", use_container_width=True):
                # Validation
                effective_objective = st.session_state.smart_objective or objective
                effective_audience = st.session_state.audience_profile or target_audience
                missing = []
                if not full_background.strip():
                    missing.append("Background/Context")
                if not effective_objective.strip():
                    missing.append("Objective")
                if not effective_audience.strip():
                    missing.append("Target Audience")
                if missing:
                    st.error(f"Please fill in all required fields: {', '.join(missing)}")
                else:
                    brief["background"] = full_background
                    brief["objective"] = effective_objective
                    brief["target_audience"] = effective_audience
                    st.session_state.current_brief = brief
                    st.session_state.wizard_step = 3
                    st.rerun()

    # --- Step 3: Messaging ---
    elif step == 3:
        st.subheader("Step 3: Messaging")

        # --- Key Insight ---
        st.markdown("#### Key Insight")
        key_insight = st.text_area(
            "Key Insight",
            value=brief.get("key_insight", ""),
            height=80,
            placeholder='What frustration or desire does your audience feel? e.g., "Sellers didn\'t start their business to write product descriptions..."',
            help=(
                "The key insight is the emotional reason your audience will care about this campaign. "
                "It's not a data point — it's a **feeling**.\n\n"
                "**Example for AI Magic Listing:**\n"
                '*"Sellers spend hours crafting listings when all they really want is to hit publish '
                'and start selling. Every minute spent writing descriptions feels like a minute stolen '
                'from growing their business."*\n\n'
                "Click **Extract Insight** to auto-generate from your background and audience."
            ),
            label_visibility="collapsed",
        )

        col1, col2 = st.columns(2)
        with col1:
            if st.button("Extract Insight", key="btn_extract_insight", use_container_width=True):
                if generator and not st.session_state.demo_mode:
                    with st.spinner("Extracting key insight..."):
                        try:
                            insight = generator.extract_insight(
                                brief.get("background", ""),
                                brief.get("target_audience", ""),
                            )
                            st.session_state.ai_insight = insight
                        except Exception as e:
                            st.error(f"Error: {e}")
                else:
                    st.session_state.ai_insight = (
                        "Professional sellers spend 15+ hours per week on listing creation -- "
                        "they want more time selling, not typing descriptions. The tedium of "
                        "manual listing isn't just annoying, it's actively costing them money."
                    )
                st.rerun()
        with col2:
            if st.button("Proofread", key="btn_proof_insight", use_container_width=True):
                if not key_insight.strip():
                    st.warning("Please type an insight first before proofreading.")
                elif generator and not st.session_state.demo_mode:
                    with st.spinner("Proofreading..."):
                        try:
                            proofed = generator.proofread_text(key_insight)
                            st.session_state.ai_insight = proofed
                        except Exception as e:
                            st.error(f"Error: {e}")
                    st.rerun()
                else:
                    st.session_state.ai_insight = key_insight.strip() + " [Proofread: looks good -- clear and emotionally resonant.]"
                    st.rerun()

        if st.session_state.get("ai_insight"):
            col_ins_info, col_ins_close = st.columns([11, 1])
            with col_ins_info:
                st.info(st.session_state.ai_insight)
            with col_ins_close:
                st.write("")
                if st.button("Close X", key="close_insight"):
                    st.session_state.ai_insight = None
                    st.rerun()
            if st.button("Use This", key="use_insight"):
                brief["key_insight"] = st.session_state.ai_insight
                st.session_state.ai_insight = None
                st.session_state.current_brief = brief
                st.rerun()

        st.divider()

        # --- Short Positioning ---
        st.markdown("#### Positioning Statement -- Short (25 words)")
        positioning_short = st.text_input(
            "Short Positioning",
            value=brief.get("positioning_short", ""),
            placeholder="e.g., AI-powered listings that sell faster with zero effort.",
            label_visibility="collapsed",
        )

        # --- Detailed Positioning ---
        st.markdown("#### Positioning Statement -- Detailed (100 words)")
        positioning_detailed = st.text_area(
            "Detailed Positioning",
            value=brief.get("positioning_detailed", ""),
            height=100,
            placeholder="For [target audience] who [need], [product] is the [category] that [key benefit]...",
            label_visibility="collapsed",
        )

        col1, col2 = st.columns(2)
        with col1:
            if st.button("Generate Positioning", key="btn_gen_pos", use_container_width=True):
                if generator and not st.session_state.demo_mode:
                    with st.spinner("Generating positioning statements..."):
                        try:
                            pos = generator.generate_positioning(brief)
                            st.session_state.ai_positioning = pos
                        except Exception as e:
                            st.error(f"Error: {e}")
                else:
                    st.session_state.ai_positioning = {
                        "positioning_short": "eBay Magical Listing turns a single photo into a complete, optimized listing in seconds.",
                        "positioning_detailed": (
                            "For professional sellers who are tired of spending hours on manual listings, "
                            "eBay Magical Listing is an AI-powered tool that transforms a single product "
                            "photo into a fully optimized listing with title, description, item specifics, "
                            "and pricing -- saving up to 80% of listing time while improving listing "
                            "quality and conversion rates compared to manual creation."
                        ),
                    }
                st.rerun()
        with col2:
            if st.button("Proofread", key="btn_proof_pos", use_container_width=True):
                combined = f"{positioning_short}\n\n{positioning_detailed}".strip()
                if not combined.strip():
                    st.warning("Please type positioning text first before proofreading.")
                elif generator and not st.session_state.demo_mode:
                    with st.spinner("Proofreading..."):
                        try:
                            proofed = generator.proofread_text(combined)
                            st.session_state.ai_positioning = {
                                "positioning_short": proofed.split("\n\n")[0] if "\n\n" in proofed else proofed,
                                "positioning_detailed": proofed.split("\n\n", 1)[1] if "\n\n" in proofed else "",
                            }
                        except Exception as e:
                            st.error(f"Error: {e}")
                    st.rerun()
                else:
                    st.session_state.ai_positioning = {
                        "positioning_short": positioning_short + " [Proofread: concise and clear.]",
                        "positioning_detailed": positioning_detailed + " [Proofread: well-structured.]",
                    }
                    st.rerun()

        if st.session_state.get("ai_positioning"):
            pos = st.session_state.ai_positioning
            col_pos_info, col_pos_close = st.columns([11, 1])
            with col_pos_info:
                st.info(f"**Short:** {pos.get('positioning_short', pos.get('short', ''))}")
                st.info(f"**Detailed:** {pos.get('positioning_detailed', pos.get('detailed', ''))}")
            with col_pos_close:
                st.write("")
                if st.button("Close X", key="close_pos"):
                    st.session_state.ai_positioning = None
                    st.rerun()
            if st.button("Use This", key="use_pos"):
                brief["positioning_short"] = pos.get("positioning_short", pos.get("short", ""))
                brief["positioning_detailed"] = pos.get("positioning_detailed", pos.get("detailed", ""))
                st.session_state.ai_positioning = None
                st.session_state.current_brief = brief
                st.rerun()

        st.divider()

        # --- Key Messages ---
        st.markdown("#### Key Messages (3-5 ranked)")
        key_messages = st.text_area(
            "Key Messages",
            value=brief.get("key_messages_text", "\n".join(brief.get("key_messages", [])) if isinstance(brief.get("key_messages"), list) else brief.get("key_messages", "")),
            height=120,
            placeholder="Enter one message per line, ranked by priority:\n1. Most important message...\n2. Second message...\n3. Third message...",
            label_visibility="collapsed",
        )

        col1, col2 = st.columns(2)
        with col1:
            if st.button("Generate Key Messages", key="btn_gen_msgs", use_container_width=True):
                if generator and not st.session_state.demo_mode:
                    with st.spinner("Generating key messages..."):
                        try:
                            msgs = generator.generate_key_messages(brief)
                            st.session_state.ai_messages = msgs
                        except Exception as e:
                            st.error(f"Error: {e}")
                else:
                    st.session_state.ai_messages = [
                        "List in seconds, not hours -- snap a photo and let AI do the rest.",
                        "AI-optimized titles and descriptions that buyers actually search for.",
                        "Over 10 million sellers already use Magical Listing -- join them.",
                        "Focus on what you do best: sourcing and selling.",
                        "Better listings mean better visibility, more traffic, and more sales.",
                    ]
                st.rerun()
        with col2:
            if st.button("Proofread", key="btn_proof_msgs", use_container_width=True):
                if not key_messages.strip():
                    st.warning("Please type key messages first before proofreading.")
                elif generator and not st.session_state.demo_mode:
                    with st.spinner("Proofreading..."):
                        try:
                            proofed = generator.proofread_text(key_messages)
                            st.session_state.ai_messages = proofed
                        except Exception as e:
                            st.error(f"Error: {e}")
                    st.rerun()
                else:
                    st.session_state.ai_messages = key_messages.strip() + "\n[Proofread: messages are clear and prioritized.]"
                    st.rerun()

        if st.session_state.get("ai_messages"):
            msgs = st.session_state.ai_messages
            col_msgs_info, col_msgs_close = st.columns([11, 1])
            with col_msgs_info:
                if isinstance(msgs, list):
                    for i, m in enumerate(msgs, 1):
                        st.info(f"**{i}.** {m}")
                else:
                    st.info(msgs)
            with col_msgs_close:
                st.write("")
                if st.button("Close X", key="close_msgs"):
                    st.session_state.ai_messages = None
                    st.rerun()
            if st.button("Use These", key="use_msgs"):
                if isinstance(msgs, list):
                    brief["key_messages"] = msgs
                    brief["key_messages_text"] = "\n".join(f"{i+1}. {m}" for i, m in enumerate(msgs))
                else:
                    brief["key_messages_text"] = msgs
                st.session_state.ai_messages = None
                st.session_state.current_brief = brief
                st.rerun()

        st.divider()

        # --- Single-Minded Proposition (SMP) ---
        st.markdown("#### Single-Minded Proposition (SMP)")
        st.caption("The ONE thing this campaign communicates. If it says two things, it's not single-minded.")
        smp_text = st.text_input(
            "SMP",
            value=brief.get("smp", ""),
            placeholder="e.g., AI Magic Listing cuts your listing time in half.",
            label_visibility="collapsed",
        )

        if st.button("Generate SMP", key="btn_gen_smp", use_container_width=True):
            if generator and not st.session_state.demo_mode:
                with st.spinner("Generating Single-Minded Proposition..."):
                    try:
                        smp = generator.generate_smp(brief)
                        st.session_state.ai_smp = smp
                    except Exception as e:
                        st.error(f"Error: {e}")
            else:
                st.session_state.ai_smp = {
                    "smp": "One photo. One tap. One perfect listing.",
                    "pass": True,
                    "reason": "No 'and/also' language detected. Proposition is genuinely single-minded.",
                }
            st.rerun()

        if st.session_state.get("ai_smp"):
            smp_data = st.session_state.ai_smp
            smp_val = smp_data.get("smp", "")
            qc = smp_data.get("quality_check", smp_data.get("pass", ""))
            qr = smp_data.get("quality_reason", smp_data.get("reason", ""))

            col_smp_info, col_smp_close = st.columns([11, 1])
            with col_smp_info:
                if qc == "pass" or qc is True:
                    st.success(f"**SMP:** {smp_val}  \n**Quality Check:** PASS -- {qr}")
                else:
                    st.error(f"**SMP:** {smp_val}  \n**Quality Check:** FAIL -- {qr}")
            with col_smp_close:
                st.write("")
                if st.button("Close X", key="close_smp"):
                    st.session_state.ai_smp = None
                    st.rerun()

            if st.button("Use This", key="use_smp"):
                brief["smp"] = smp_val
                brief["smp_pass"] = (qc == "pass" or qc is True)
                st.session_state.ai_smp = None
                st.session_state.current_brief = brief
                st.rerun()

        # --- Navigation ---
        st.write("")
        col_back, col_next = st.columns(2)
        with col_back:
            if st.button("<<  Back", use_container_width=True, key="back3"):
                st.session_state.wizard_step = 2
                st.rerun()
        with col_next:
            if st.button("Next  >>", type="primary", use_container_width=True, key="next3"):
                # Save whatever the user has typed or accepted
                brief["key_insight"] = key_insight or brief.get("key_insight", "")
                brief["positioning_short"] = positioning_short or brief.get("positioning_short", "")
                brief["positioning_detailed"] = positioning_detailed or brief.get("positioning_detailed", "")
                if key_messages.strip():
                    brief["key_messages_text"] = key_messages
                    brief["key_messages"] = [m.strip().lstrip("0123456789.)-) ") for m in key_messages.strip().split("\n") if m.strip()]
                brief["smp"] = smp_text or brief.get("smp", "")

                # Validation
                missing = []
                if not brief.get("key_insight") and not key_insight.strip():
                    missing.append("Key Insight")
                if not brief.get("positioning_short") and not positioning_short.strip():
                    missing.append("Positioning (Short)")
                if not brief.get("key_messages") and not key_messages.strip():
                    missing.append("Key Messages")
                if not brief.get("smp") and not smp_text.strip():
                    missing.append("SMP")

                if missing:
                    st.error(f"Please fill in: {', '.join(missing)}")
                else:
                    st.session_state.current_brief = brief
                    st.session_state.wizard_step = 4
                    st.rerun()

    # --- Step 4: Execution ---
    elif step == 4:
        st.subheader("Step 4: Execution Plan")

        # --- Channel Plan ---
        st.markdown("#### Channel Plan")
        channel_plan_text = st.text_area(
            "Channel Plan",
            value=brief.get("channel_plan_text", ""),
            height=120,
            placeholder="Describe your channel strategy, or click Generate to create one:\ne.g., Email (30%) - drip campaign to sellers\nIn-App (25%) - contextual notifications\nSocial (20%) - demo videos on TikTok/Instagram",
            label_visibility="collapsed",
        )

        col1, col2 = st.columns(2)
        with col1:
            if st.button("Generate Channel Plan", key="btn_gen_channels", use_container_width=True):
                if generator and not st.session_state.demo_mode:
                    with st.spinner("Generating channel plan..."):
                        try:
                            channels = generator.generate_channel_plan(brief)
                            st.session_state.ai_channels = channels
                        except Exception as e:
                            st.error(f"Error: {e}")
                else:
                    st.session_state.ai_channels = [
                        {"channel": "Email", "tactic": "Segmented drip campaign to pro sellers", "rationale": "High engagement with seller comms", "budget_pct": 25},
                        {"channel": "In-App Notifications", "tactic": "Contextual prompts during listing flow", "rationale": "Catch sellers at point of need", "budget_pct": 15},
                        {"channel": "Seller Hub Banners", "tactic": "Feature spotlight banner with demo video", "rationale": "Primary seller workspace", "budget_pct": 10},
                        {"channel": "Social Media", "tactic": "Before/after listing demos on Instagram & TikTok", "rationale": "Visual impact, shareable content", "budget_pct": 20},
                        {"channel": "YouTube", "tactic": "Tutorial series: Magical Listing mastery", "rationale": "Educational content for adoption", "budget_pct": 15},
                        {"channel": "Paid Search", "tactic": "Target 'sell on eBay' and competitor queries", "rationale": "Capture high-intent sellers", "budget_pct": 15},
                    ]
                st.rerun()
        with col2:
            if st.button("Proofread", key="btn_proof_channels", use_container_width=True):
                if not channel_plan_text.strip():
                    st.warning("Write some channel plan text first, then click Proofread.")
                else:
                    if generator and not st.session_state.demo_mode:
                        with st.spinner("Proofreading..."):
                            try:
                                result = generator.proofread_text(channel_plan_text)
                                st.session_state.ai_channels_proofed = result
                            except Exception as e:
                                st.error(f"Error: {e}")
                    else:
                        st.session_state.ai_channels_proofed = (
                            channel_plan_text.rstrip(".") + ". "
                            "[Proofread] Channel plan refined for clarity, consistency, "
                            "and professional tone."
                        )
                    st.rerun()

        if st.session_state.get("ai_channels_proofed"):
            col_ch_info, col_ch_close = st.columns([11, 1])
            with col_ch_info:
                st.info(st.session_state.ai_channels_proofed)
            with col_ch_close:
                st.write("")
                if st.button("Close X", key="close_channels_proofed"):
                    st.session_state.ai_channels_proofed = None
                    st.rerun()
            if st.button("Use This", key="use_channels_proofed"):
                brief["channel_plan_text"] = st.session_state.ai_channels_proofed
                st.session_state.ai_channels_proofed = None
                st.session_state.current_brief = brief
                st.rerun()

        if st.session_state.get("ai_channels"):
            channel_display = []
            for ch in st.session_state.ai_channels:
                channel_display.append({
                    "Channel": ch.get("channel", ""),
                    "Tactic": ch.get("tactic", ""),
                    "Rationale": ch.get("rationale", ""),
                    "Budget %": ch.get("budget_pct", ch.get("budget_percent", "")),
                })
            df_channels = pd.DataFrame(channel_display)
            st.dataframe(df_channels, use_container_width=True, hide_index=True)
            if st.button("Use This", key="use_channels"):
                brief["channel_plan"] = st.session_state.ai_channels
                brief["channel_plan_text"] = "\n".join(
                    f"{ch.get('channel', '')}: {ch.get('tactic', '')} ({ch.get('budget_pct', '')}%)"
                    for ch in st.session_state.ai_channels
                )
                st.session_state.ai_channels = None
                st.session_state.current_brief = brief
                st.rerun()

        st.divider()

        # --- Content Deliverables ---
        st.markdown("#### Content Deliverables")
        deliverables_text = st.text_area(
            "Deliverables",
            value=brief.get("deliverables_text", ""),
            height=120,
            placeholder="List your content deliverables, or click Generate:\ne.g., 5x email templates - HTML responsive\n3x social videos - 15s/30s/60s\n1x blog post - 800 words, SEO optimized",
            label_visibility="collapsed",
        )

        col1, col2 = st.columns(2)
        with col1:
            if st.button("Generate Deliverables", key="btn_gen_deliverables", use_container_width=True):
                if generator and not st.session_state.demo_mode:
                    with st.spinner("Generating deliverables..."):
                        try:
                            deliverables = generator.generate_deliverables(brief)
                            st.session_state.ai_deliverables = deliverables
                        except Exception as e:
                            st.error(f"Error generating deliverables: {e}")
                else:
                    st.session_state.ai_deliverables = [
                        {"asset": "Email Templates (x5)", "spec": "HTML responsive, 200 words max each", "owner": "PMM"},
                        {"asset": "Social Video Clips (x3)", "spec": "15s/30s/60s, 1080x1080 and 9:16", "owner": "Video"},
                        {"asset": "YouTube Tutorial", "spec": "5 min, 1920x1080", "owner": "Video"},
                        {"asset": "Seller Hub Banners (x2)", "spec": "1200x400px", "owner": "Design"},
                        {"asset": "Landing Page", "spec": "Responsive web with interactive demo", "owner": "Web"},
                        {"asset": "In-App Notifications (x10)", "spec": "100 chars max each", "owner": "Product"},
                        {"asset": "Blog Post", "spec": "800 words, SEO optimized", "owner": "Content"},
                    ]
                st.rerun()
        with col2:
            if st.button("Proofread", key="btn_proof_deliverables", use_container_width=True):
                if not deliverables_text.strip():
                    st.warning("Write some deliverables text first, then click Proofread.")
                else:
                    if generator and not st.session_state.demo_mode:
                        with st.spinner("Proofreading..."):
                            try:
                                result = generator.proofread_text(deliverables_text)
                                st.session_state.ai_deliverables_proofed = result
                            except Exception as e:
                                st.error(f"Error: {e}")
                    else:
                        st.session_state.ai_deliverables_proofed = (
                            deliverables_text.rstrip(".") + ". "
                            "[Proofread] Deliverables refined for clarity and consistent formatting."
                        )
                    st.rerun()

        if st.session_state.get("ai_deliverables_proofed"):
            col_del_info, col_del_close = st.columns([11, 1])
            with col_del_info:
                st.info(st.session_state.ai_deliverables_proofed)
            with col_del_close:
                st.write("")
                if st.button("Close X", key="close_deliverables_proofed"):
                    st.session_state.ai_deliverables_proofed = None
                    st.rerun()
            if st.button("Use This", key="use_deliverables_proofed"):
                brief["deliverables_text"] = st.session_state.ai_deliverables_proofed
                st.session_state.ai_deliverables_proofed = None
                st.session_state.current_brief = brief
                st.rerun()

        if st.session_state.get("ai_deliverables"):
            del_data = st.session_state.ai_deliverables
            if del_data and isinstance(del_data[0], dict):
                df_del = pd.DataFrame(del_data)
                df_del.columns = [c.replace("_", " ").title() for c in df_del.columns]
                st.dataframe(df_del, use_container_width=True, hide_index=True)
            else:
                for d in del_data:
                    st.write(f"- {d}")
            if st.button("Use This", key="use_deliverables"):
                brief["deliverables"] = st.session_state.ai_deliverables
                if isinstance(del_data[0], dict):
                    brief["deliverables_text"] = "\n".join(
                        f"{d.get('asset', '')}: {d.get('spec', '')} — {d.get('owner', '')}"
                        for d in del_data
                    )
                st.session_state.ai_deliverables = None
                st.session_state.current_brief = brief
                st.rerun()

        st.divider()

        # --- Timeline ---
        st.markdown("#### Campaign Timeline")
        timeline_text = st.text_area(
            "Timeline",
            value=brief.get("timeline_text", ""),
            height=120,
            placeholder="Describe your timeline, or click Generate:\ne.g., Awareness (Weeks 1-2): teaser emails, social posts\nLaunch (Weeks 3-4): full campaign, in-app prompts\nSustain (Weeks 5-8): success stories, A/B testing\nOptimize (Weeks 9+): analyze, scale winners",
            label_visibility="collapsed",
        )

        col1, col2 = st.columns(2)
        with col1:
            if st.button("Generate Timeline", key="btn_gen_timeline", use_container_width=True):
                if generator and not st.session_state.demo_mode:
                    with st.spinner("Generating timeline..."):
                        try:
                            timeline = generator.generate_timeline(brief)
                            st.session_state.ai_timeline = timeline
                        except Exception as e:
                            st.error(f"Error generating timeline: {e}")
                else:
                    st.session_state.ai_timeline = [
                        {"phase": "Awareness", "duration": "Weeks 1-2", "actions": "Teaser emails, social posts, blog announcement, seller community seeding"},
                        {"phase": "Launch", "duration": "Weeks 3-4", "actions": "Full campaign launch, in-app prompts, YouTube video, paid search activation"},
                        {"phase": "Sustain", "duration": "Weeks 5-8", "actions": "Success stories, re-engagement emails, A/B testing, community engagement"},
                        {"phase": "Optimize", "duration": "Weeks 9-12", "actions": "Analyze results, scale winning channels, cut underperformers, report"},
                    ]
                st.rerun()
        with col2:
            if st.button("Proofread", key="btn_proof_timeline", use_container_width=True):
                if not timeline_text.strip():
                    st.warning("Write some timeline text first, then click Proofread.")
                else:
                    if generator and not st.session_state.demo_mode:
                        with st.spinner("Proofreading..."):
                            try:
                                result = generator.proofread_text(timeline_text)
                                st.session_state.ai_timeline_proofed = result
                            except Exception as e:
                                st.error(f"Error: {e}")
                    else:
                        st.session_state.ai_timeline_proofed = (
                            timeline_text.rstrip(".") + ". "
                            "[Proofread] Timeline refined for clarity and consistent phase structure."
                        )
                    st.rerun()

        if st.session_state.get("ai_timeline_proofed"):
            col_tl_info, col_tl_close = st.columns([11, 1])
            with col_tl_info:
                st.info(st.session_state.ai_timeline_proofed)
            with col_tl_close:
                st.write("")
                if st.button("Close X", key="close_timeline_proofed"):
                    st.session_state.ai_timeline_proofed = None
                    st.rerun()
            if st.button("Use This", key="use_timeline_proofed"):
                brief["timeline_text"] = st.session_state.ai_timeline_proofed
                st.session_state.ai_timeline_proofed = None
                st.session_state.current_brief = brief
                st.rerun()

        if st.session_state.get("ai_timeline"):
            timeline_display = []
            for t in st.session_state.ai_timeline:
                actions = t.get("actions", t.get("activities", ""))
                if isinstance(actions, list):
                    actions = ", ".join(actions)
                timeline_display.append({
                    "Phase": t.get("phase", ""),
                    "Duration": t.get("duration", ""),
                    "Activities": actions,
                })
            df_timeline = pd.DataFrame(timeline_display)
            st.dataframe(df_timeline, use_container_width=True, hide_index=True)
            if st.button("Use This", key="use_timeline"):
                brief["timeline"] = st.session_state.ai_timeline
                brief["timeline_text"] = "\n".join(
                    f"{t.get('phase', '')} ({t.get('duration', '')}): {', '.join(t.get('actions', t.get('activities', []))) if isinstance(t.get('actions', t.get('activities', '')), list) else t.get('actions', t.get('activities', ''))}"
                    for t in st.session_state.ai_timeline
                )
                st.session_state.ai_timeline = None
                st.session_state.current_brief = brief
                st.rerun()

        st.divider()

        # --- Budget (optional) ---
        st.markdown("#### Budget (optional)")
        budget_input = st.text_input(
            "Total Budget",
            value=brief.get("budget", ""),
            placeholder="e.g., $150,000",
        )

        # Show budget pie chart if we have channel plan with budget %
        if budget_input and brief.get("channel_plan"):
            channels = brief["channel_plan"]
            labels = [ch.get("channel", "") for ch in channels]
            values = [ch.get("budget_pct", 0) for ch in channels]
            if any(values):
                fig = px.pie(
                    names=labels,
                    values=values,
                    title="Budget Allocation by Channel",
                    color_discrete_sequence=[PRIMARY_BLUE, ACCENT_RED, ACCENT_YELLOW, ACCENT_GREEN, "#9B59B6", "#1ABC9C"],
                )
                fig.update_layout(
                    plot_bgcolor="rgba(0,0,0,0)",
                    paper_bgcolor="rgba(0,0,0,0)",
                    font=dict(size=12),
                    height=350,
                )
                st.plotly_chart(fig, use_container_width=True)

        # --- Navigation ---
        st.write("")
        col_back, col_next = st.columns(2)
        with col_back:
            if st.button("<<  Back", use_container_width=True, key="back4"):
                st.session_state.wizard_step = 3
                st.rerun()
        with col_next:
            if st.button("Next  >>", type="primary", use_container_width=True, key="next4"):
                # Save text fields
                if channel_plan_text.strip():
                    brief["channel_plan_text"] = channel_plan_text
                if deliverables_text.strip():
                    brief["deliverables_text"] = deliverables_text
                if timeline_text.strip():
                    brief["timeline_text"] = timeline_text
                if budget_input:
                    brief["budget"] = budget_input

                # Validation — need at least channel plan and timeline
                missing = []
                if not brief.get("channel_plan") and not channel_plan_text.strip():
                    missing.append("Channel Plan")
                if not brief.get("timeline") and not timeline_text.strip():
                    missing.append("Timeline")

                if missing:
                    st.error(f"Please fill in: {', '.join(missing)}")
                else:
                    st.session_state.current_brief = brief
                    st.session_state.wizard_step = 5
                    st.rerun()

    # --- Step 5: Governance ---
    elif step == 5:
        st.subheader("Step 5: Governance")

        # --- Success Metrics / KPIs ---
        kpis_text = st.text_area(
            "Success Metrics / KPIs",
            value=brief.get("kpis_text", ""),
            height=120,
            placeholder="What metrics will define success? Include targets and how you'll measure them.",
        )
        if kpis_text != brief.get("kpis_text", ""):
            brief["kpis_text"] = kpis_text
            st.session_state.current_brief = brief

        if st.button("Generate KPIs", use_container_width=True):
            if generator and not st.session_state.demo_mode:
                with st.spinner("Generating KPIs..."):
                    try:
                        kpis = generator.generate_kpis(brief)
                        st.session_state.generated_kpis = kpis
                    except Exception as e:
                        st.error(f"Error: {e}")
            else:
                st.session_state.generated_kpis = [
                    {"metric": "Magical Listing Adoption Rate", "target": "+40% from baseline", "measurement": "Product analytics dashboard"},
                    {"metric": "Listing Creation Time", "target": "-60% reduction", "measurement": "Session analytics"},
                    {"metric": "Email Open Rate", "target": ">25%", "measurement": "Email marketing platform"},
                    {"metric": "Video Completion Rate", "target": ">50%", "measurement": "YouTube / Social analytics"},
                    {"metric": "New Seller Signups (attributed)", "target": "+15%", "measurement": "Multi-touch attribution model"},
                    {"metric": "Listing Quality Score", "target": "+20% improvement", "measurement": "Internal listing quality index"},
                ]
            st.rerun()

        if st.session_state.generated_kpis:
            kpi_display = []
            for k in st.session_state.generated_kpis:
                kpi_display.append({
                    "Metric": k.get("metric", ""),
                    "Target": k.get("target", ""),
                    "Measurement": k.get("measurement", k.get("measurement_method", "")),
                })
            df_kpis = pd.DataFrame(kpi_display)
            col_kpi_info, col_kpi_close = st.columns([11, 1])
            with col_kpi_info:
                st.dataframe(df_kpis, use_container_width=True, hide_index=True)
            with col_kpi_close:
                st.write("")
                if st.button("Close X", key="close_kpis"):
                    st.session_state.generated_kpis = None
                    st.rerun()
            if st.button("Use These KPIs", key="use_kpis"):
                brief["kpis"] = st.session_state.generated_kpis
                # Also populate the text box with a readable version
                kpi_lines = []
                for k in st.session_state.generated_kpis:
                    kpi_lines.append(f"- {k.get('metric', '')}: {k.get('target', '')} (measured via {k.get('measurement', k.get('measurement_method', ''))})")
                brief["kpis_text"] = "\n".join(kpi_lines)
                st.session_state.generated_kpis = None
                st.session_state.current_brief = brief
                st.rerun()

        st.divider()

        # --- RACI Matrix ---
        raci_text = st.text_area(
            "RACI Matrix",
            value=brief.get("raci_text", ""),
            height=120,
            placeholder="Who is Responsible, Accountable, Consulted, and Informed for each workstream?",
        )
        if raci_text != brief.get("raci_text", ""):
            brief["raci_text"] = raci_text
            st.session_state.current_brief = brief

        if st.button("Generate RACI", use_container_width=True):
            if generator and not st.session_state.demo_mode:
                with st.spinner("Generating RACI matrix..."):
                    try:
                        raci = generator.generate_raci_matrix(brief.get("brief_type", "product_launch"))
                        st.session_state.generated_raci = raci
                    except Exception as e:
                        st.error(f"Error: {e}")
            else:
                st.session_state.generated_raci = [
                    {"task": "Campaign Strategy & Brief", "responsible": "PMM Lead", "accountable": "VP Marketing", "consulted": "Product, Data Science", "informed": "Executive Team"},
                    {"task": "Creative Asset Development", "responsible": "Creative Team", "accountable": "PMM Lead", "consulted": "Brand Guidelines", "informed": "PMM"},
                    {"task": "Email Campaign Execution", "responsible": "Email Marketing", "accountable": "PMM Lead", "consulted": "CRM Team", "informed": "Analytics"},
                    {"task": "Social Media Campaign", "responsible": "Social Team", "accountable": "PMM Lead", "consulted": "Creative, Legal", "informed": "PR"},
                    {"task": "Paid Search Management", "responsible": "Performance Marketing", "accountable": "PMM Lead", "consulted": "Finance", "informed": "Analytics"},
                    {"task": "Performance Reporting", "responsible": "Analytics Team", "accountable": "PMM Lead", "consulted": "Data Science", "informed": "Executive Team"},
                ]
            st.rerun()

        if st.session_state.generated_raci:
            raci_display = []
            for r in st.session_state.generated_raci:
                raci_display.append({
                    "Task": r.get("task", r.get("role", "")),
                    "Responsible": r.get("responsible", r.get("R", "")),
                    "Accountable": r.get("accountable", r.get("A", "")),
                    "Consulted": r.get("consulted", r.get("C", "")),
                    "Informed": r.get("informed", r.get("I", "")),
                })
            df_raci = pd.DataFrame(raci_display)
            col_raci_info, col_raci_close = st.columns([11, 1])
            with col_raci_info:
                st.dataframe(df_raci, use_container_width=True, hide_index=True)
            with col_raci_close:
                st.write("")
                if st.button("Close X", key="close_raci"):
                    st.session_state.generated_raci = None
                    st.rerun()
            if st.button("Use This RACI", key="use_raci"):
                brief["raci"] = st.session_state.generated_raci
                # Also populate the text box with a readable version
                raci_lines = []
                for r in st.session_state.generated_raci:
                    raci_lines.append(f"- {r.get('task', r.get('role', ''))}: R={r.get('responsible', '')} | A={r.get('accountable', '')} | C={r.get('consulted', '')} | I={r.get('informed', '')}")
                brief["raci_text"] = "\n".join(raci_lines)
                st.session_state.generated_raci = None
                st.session_state.current_brief = brief
                st.rerun()

        st.write("")
        col_back, col_next = st.columns(2)
        with col_back:
            if st.button("<<  Back", use_container_width=True, key="back5"):
                st.session_state.wizard_step = 4
                st.rerun()
        with col_next:
            if st.button("Next  >>", type="primary", use_container_width=True, key="next5"):
                missing = []
                if not st.session_state.generated_kpis and not brief.get("kpis") and not brief.get("kpis_text", "").strip():
                    missing.append("KPIs (enter text or click 'Generate KPIs')")
                if not st.session_state.generated_raci and not brief.get("raci") and not brief.get("raci_text", "").strip():
                    missing.append("RACI Matrix (enter text or click 'Generate RACI')")
                if missing:
                    st.error(f"Please complete the following before proceeding: {', '.join(missing)}")
                else:
                    if st.session_state.generated_kpis:
                        brief["kpis"] = st.session_state.generated_kpis
                    if st.session_state.generated_raci:
                        brief["raci"] = st.session_state.generated_raci
                    st.session_state.current_brief = brief
                    st.session_state.wizard_step = 6
                    st.rerun()

    # --- Step 6: Review & Export ---
    elif step == 6:
        st.subheader("Step 6: Review & Export")

        brief["created"] = brief.get("created", datetime.now().isoformat())
        md = brief_to_markdown(brief)
        word_count = len(md.split())

        with st.expander("Full Brief Preview", expanded=True):
            st.markdown(md)

        st.write("")
        col_dl, col_save = st.columns(2)

        with col_dl:
            # Generate Word document
            try:
                from docx import Document as DocxDocument
                from docx.shared import Pt, Inches, RGBColor
                from docx.enum.text import WD_ALIGN_PARAGRAPH

                doc = DocxDocument()

                # Title
                title = doc.add_heading(f"Campaign Brief: {brief.get('campaign_name', 'Untitled')}", level=0)
                meta = doc.add_paragraph()
                meta.add_run(f"Type: {brief.get('brief_type', 'N/A')}  |  Tier: {brief.get('launch_tier', 'N/A')}  |  Date: {brief.get('created', '')[:10]}").font.size = Pt(10)

                def _docx_section(title, content):
                    doc.add_heading(title, level=2)
                    doc.add_paragraph(str(content) if content else "Not specified")

                _docx_section("Background & Context", brief.get("background"))
                _docx_section("Objective", brief.get("objective"))
                _docx_section("Target Audience", brief.get("target_audience"))
                _docx_section("Key Insight", brief.get("key_insight"))

                if brief.get("positioning_short"):
                    doc.add_heading("Positioning Statement", level=2)
                    p = doc.add_paragraph()
                    p.add_run("Short: ").bold = True
                    p.add_run(brief["positioning_short"])
                    if brief.get("positioning_detailed"):
                        p2 = doc.add_paragraph()
                        p2.add_run("Detailed: ").bold = True
                        p2.add_run(brief["positioning_detailed"])

                if brief.get("key_messages"):
                    doc.add_heading("Key Messages", level=2)
                    msgs = brief["key_messages"]
                    if isinstance(msgs, list):
                        for i, m in enumerate(msgs, 1):
                            doc.add_paragraph(f"{m}", style="List Number")
                    else:
                        doc.add_paragraph(str(msgs))

                if brief.get("smp"):
                    doc.add_heading("Single-Minded Proposition (SMP)", level=2)
                    p = doc.add_paragraph()
                    p.add_run(brief["smp"]).bold = True
                    if brief.get("smp_pass") is not None:
                        doc.add_paragraph(f"Quality Check: {'PASS' if brief['smp_pass'] else 'FAIL'}")

                if brief.get("channel_plan"):
                    doc.add_heading("Channel Plan", level=2)
                    table = doc.add_table(rows=1, cols=4)
                    table.style = "Light Grid Accent 1"
                    hdr = table.rows[0].cells
                    hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = "Channel", "Tactic", "Rationale", "Budget %"
                    for ch in brief["channel_plan"]:
                        row = table.add_row().cells
                        row[0].text = str(ch.get("channel", ""))
                        row[1].text = str(ch.get("tactic", ""))
                        row[2].text = str(ch.get("rationale", ""))
                        row[3].text = str(ch.get("budget_pct", ""))

                if brief.get("deliverables"):
                    doc.add_heading("Content Deliverables", level=2)
                    for d in brief["deliverables"]:
                        if isinstance(d, dict):
                            p = doc.add_paragraph(style="List Bullet")
                            p.add_run(f"{d.get('asset', '')}").bold = True
                            p.add_run(f" — {d.get('spec', '')} (Owner: {d.get('owner', '')})")
                        else:
                            doc.add_paragraph(str(d), style="List Bullet")

                if brief.get("timeline"):
                    doc.add_heading("Timeline", level=2)
                    table = doc.add_table(rows=1, cols=3)
                    table.style = "Light Grid Accent 1"
                    hdr = table.rows[0].cells
                    hdr[0].text, hdr[1].text, hdr[2].text = "Phase", "Duration", "Activities"
                    for t in brief["timeline"]:
                        row = table.add_row().cells
                        row[0].text = str(t.get("phase", ""))
                        row[1].text = str(t.get("duration", ""))
                        acts = t.get("actions", t.get("activities", ""))
                        row[2].text = ", ".join(acts) if isinstance(acts, list) else str(acts)

                if brief.get("budget"):
                    doc.add_heading("Budget", level=2)
                    doc.add_paragraph(f"Total: {brief['budget']}")

                if brief.get("kpis"):
                    doc.add_heading("Success Metrics / KPIs", level=2)
                    table = doc.add_table(rows=1, cols=3)
                    table.style = "Light Grid Accent 1"
                    hdr = table.rows[0].cells
                    hdr[0].text, hdr[1].text, hdr[2].text = "Metric", "Target", "Measurement"
                    for k in brief["kpis"]:
                        row = table.add_row().cells
                        row[0].text = str(k.get("metric", ""))
                        row[1].text = str(k.get("target", ""))
                        row[2].text = str(k.get("measurement", ""))

                if brief.get("raci"):
                    doc.add_heading("RACI Matrix", level=2)
                    table = doc.add_table(rows=1, cols=5)
                    table.style = "Light Grid Accent 1"
                    hdr = table.rows[0].cells
                    hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text, hdr[4].text = "Task", "Responsible", "Accountable", "Consulted", "Informed"
                    for r in brief["raci"]:
                        row = table.add_row().cells
                        row[0].text = str(r.get("task", ""))
                        row[1].text = str(r.get("responsible", ""))
                        row[2].text = str(r.get("accountable", ""))
                        row[3].text = str(r.get("consulted", ""))
                        row[4].text = str(r.get("informed", ""))

                # Save to bytes
                _docx_buffer = io.BytesIO()
                doc.save(_docx_buffer)
                _docx_bytes = _docx_buffer.getvalue()

                st.download_button(
                    "Download as Word (.docx)",
                    data=_docx_bytes,
                    file_name=f"{brief.get('campaign_name', 'brief').replace(' ', '_').lower()}_brief.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                )
            except ImportError:
                st.download_button(
                    "Download as Markdown",
                    data=md,
                    file_name=f"{brief.get('campaign_name', 'brief').replace(' ', '_').lower()}_brief.md",
                    mime="text/markdown",
                    use_container_width=True,
                )

        with col_save:
            if st.button("Save to Library", type="primary", use_container_width=True):
                brief["id"] = brief.get("id", str(uuid.uuid4()))
                brief["quality_score"] = quality_score
                # Check for duplicates
                existing_ids = {b["id"] for b in st.session_state.briefs}
                if brief["id"] not in existing_ids:
                    st.session_state.briefs.append(brief.copy())
                    st.success("Brief saved to library!")
                else:
                    # Update existing
                    for i, b in enumerate(st.session_state.briefs):
                        if b["id"] == brief["id"]:
                            st.session_state.briefs[i] = brief.copy()
                            break
                    st.success("Brief updated in library!")

        st.divider()
        st.markdown("**Next Step:** Generate a Creative Brief based on the full campaign brief above. "
                     "This creates an execution-focused document for your creative team or agency.")
        if st.button("Generate Creative Brief  >>", use_container_width=True):
            if generator and not st.session_state.demo_mode:
                with st.spinner("Generating creative brief..."):
                    try:
                        creative = CreativeEngine().generate_creative_brief(brief)
                        st.session_state.generated_creative_brief = creative
                    except Exception as e:
                        st.error(f"Error: {e}")
                st.rerun()
            else:
                st.session_state.generated_creative_brief = {
                    "project_name": f"{brief.get('campaign_name', 'Untitled')} Creative Campaign",
                    "objective": brief.get("objective", "TBD"),
                    "target_audience": brief.get("target_audience", "TBD"),
                    "single_minded_proposition": brief.get("smp", "TBD"),
                    "tone_and_manner": "Empowering, approachable, innovative. We speak as a trusted partner who makes selling easier — not as a tech company pushing features.",
                    "mandatories": [
                        "eBay brand guidelines and color palette",
                        "AI disclosure statement where applicable",
                        "Legal disclaimers on performance claims",
                        "Accessible design (WCAG 2.1 AA compliance)",
                    ],
                    "deliverables": brief.get("deliverables", ["TBD"]),
                    "inspiration_references": [
                        "Canva's 'What Will You Design Today?' campaign — simple, empowering, tool-focused",
                        "Shopify's seller success stories — authentic, relatable",
                        "Apple's product demo videos — clean, focused, aspirational",
                    ],
                    "do_nots": [
                        "Don't use overly technical AI jargon",
                        "Don't make unrealistic performance promises",
                        "Don't show competitors' platforms in comparisons",
                    ],
                }
                st.rerun()

        if st.session_state.get("generated_creative_brief"):
            cb = st.session_state.generated_creative_brief
            col_cb_info, col_cb_close = st.columns([11, 1])
            with col_cb_info:
                st.markdown("### Generated Creative Brief")
                st.markdown(_format_creative_brief(cb))
            with col_cb_close:
                st.write("")
                if st.button("Close X", key="close_creative_brief"):
                    st.session_state.generated_creative_brief = None
                    st.rerun()

            # Download Creative Brief as Word
            try:
                from docx import Document as DocxDocument
                from docx.shared import Pt

                cb_doc = DocxDocument()
                cb_doc.add_heading(f"Creative Brief: {cb.get('project_name', brief.get('campaign_name', 'Untitled'))}", level=0)

                if cb.get("objective"):
                    cb_doc.add_heading("Objective", level=2)
                    cb_doc.add_paragraph(str(cb["objective"]))
                if cb.get("target_audience"):
                    cb_doc.add_heading("Target Audience", level=2)
                    cb_doc.add_paragraph(str(cb["target_audience"]))
                if cb.get("single_minded_proposition"):
                    cb_doc.add_heading("Single-Minded Proposition", level=2)
                    p = cb_doc.add_paragraph()
                    p.add_run(str(cb["single_minded_proposition"])).italic = True
                if cb.get("tone_and_manner"):
                    cb_doc.add_heading("Tone & Manner", level=2)
                    cb_doc.add_paragraph(str(cb["tone_and_manner"]))
                if cb.get("mandatories"):
                    cb_doc.add_heading("Mandatories", level=2)
                    for m in cb["mandatories"]:
                        cb_doc.add_paragraph(str(m), style="List Bullet")
                if cb.get("deliverables"):
                    cb_doc.add_heading("Deliverables", level=2)
                    for d in cb["deliverables"]:
                        if isinstance(d, dict):
                            p = cb_doc.add_paragraph(style="List Bullet")
                            p.add_run(str(d.get("asset", ""))).bold = True
                            p.add_run(f" — {d.get('specs', d.get('spec', ''))} ({d.get('word_count', '')})")
                        else:
                            cb_doc.add_paragraph(str(d), style="List Bullet")
                if cb.get("inspiration_references"):
                    cb_doc.add_heading("Inspiration & References", level=2)
                    for r in cb["inspiration_references"]:
                        cb_doc.add_paragraph(str(r), style="List Bullet")
                if cb.get("do_nots"):
                    cb_doc.add_heading("Do Nots", level=2)
                    for d in cb["do_nots"]:
                        cb_doc.add_paragraph(str(d), style="List Bullet")

                _cb_buffer = io.BytesIO()
                cb_doc.save(_cb_buffer)
                st.download_button(
                    "Download Creative Brief as Word (.docx)",
                    data=_cb_buffer.getvalue(),
                    file_name=f"{brief.get('campaign_name', 'brief').replace(' ', '_').lower()}_creative_brief.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                )
            except ImportError:
                st.download_button(
                    "Download Creative Brief as Markdown",
                    data=_format_creative_brief(cb),
                    file_name=f"{brief.get('campaign_name', 'brief').replace(' ', '_').lower()}_creative_brief.md",
                    mime="text/markdown",
                    use_container_width=True,
                )

        st.write("")
        col_back2, col_reset = st.columns(2)
        with col_back2:
            if st.button("<<  Back", use_container_width=True, key="back6"):
                st.session_state.wizard_step = 5
                st.rerun()
        with col_reset:
            if st.button("Start New Brief", use_container_width=True):
                st.session_state.wizard_step = 1
                st.session_state.current_brief = {}
                st.session_state.generated_names = []
                st.session_state.ai_insight = None
                st.session_state.ai_positioning = None
                st.session_state.ai_messages = None
                st.session_state.ai_smp = None
                st.session_state.ai_channels = None
                st.session_state.ai_deliverables = None
                st.session_state.ai_timeline = None
                st.session_state.generated_kpis = None
                st.session_state.generated_raci = None
                st.session_state.generated_creative_brief = None
                st.session_state.audience_profile = None
                st.session_state.smart_objective = None
                st.session_state.budget_breakdown = None
                st.session_state.ai_background = None
                st.session_state.ai_generated_objective = None
                st.session_state.ai_generated_audience = None
                st.rerun()


# ---------------------------------------------------------------------------
# Tab 2: Brief Quality Checker
# ---------------------------------------------------------------------------
def render_quality_checker() -> None:
    """Render the Brief Quality Checker tab."""
    st.subheader("Brief Quality Checker")
    st.caption("Paste a brief below or upload a document to get AI-powered quality analysis.")

    brief_text = st.text_area(
        "Paste your brief here",
        value="" if not st.session_state.demo_mode else SAMPLE_BRIEF_TEXT,
        height=250,
        placeholder="Paste a campaign brief to analyze...",
    )

    uploaded_brief = st.file_uploader(
        "Or upload a brief document",
        type=["pdf", "txt", "csv", "md", "docx"],
        key="quality_doc_upload",
        help="Upload a brief as PDF, Word, text, or markdown file.",
    )
    if uploaded_brief:
        try:
            if uploaded_brief.name.endswith(".pdf"):
                try:
                    import PyPDF2
                    reader = PyPDF2.PdfReader(uploaded_brief)
                    extracted = "\n".join(page.extract_text() or "" for page in reader.pages[:20])
                except ImportError:
                    extracted = "[PDF support requires PyPDF2: pip install PyPDF2]"
            elif uploaded_brief.name.endswith(".docx"):
                try:
                    from docx import Document as DocxDocument
                    doc = DocxDocument(uploaded_brief)
                    extracted = "\n".join(p.text for p in doc.paragraphs)
                except ImportError:
                    extracted = "[DOCX support requires python-docx: pip install python-docx]"
            else:
                extracted = uploaded_brief.read().decode("utf-8", errors="ignore")

            if extracted.strip():
                brief_text = extracted
                with st.expander(f"Extracted from {uploaded_brief.name}", expanded=False):
                    st.text(extracted[:2000] + ("..." if len(extracted) > 2000 else ""))
        except Exception as e:
            st.error(f"Error reading file: {e}")

    # Brief context selectors
    col_type, col_tier = st.columns(2)
    with col_type:
        qc_brief_type = st.selectbox(
            "Brief Type",
            options=["Campaign Brief", "Creative Brief", "GTM Brief"],
            index=0,
            key="qc_brief_type",
            help="Select the type of brief you're checking so we can score against the right criteria.",
        )
    with col_tier:
        qc_tier = st.selectbox(
            "Launch Tier",
            options=["", "Tier 0", "Tier 1", "Tier 2", "Tier 3"],
            index=0,
            format_func=lambda x: "Select tier (optional)..." if x == "" else x,
            key="qc_tier",
            help="Tier 0/1 briefs are scored more strictly (all 14 sections expected). Tier 2/3 briefs are lighter.",
        )

    if st.button("Check Quality", type="primary", use_container_width=True):
        if not brief_text.strip():
            st.warning("Please paste a brief or upload a document to analyze.")
        else:
            checker = _get_quality_checker()
            # Add context for the checker
            context_prefix = f"[Brief Type: {qc_brief_type}]"
            if qc_tier:
                context_prefix += f" [{qc_tier}]"
            enriched_text = f"{context_prefix}\n\n{brief_text}"
            if checker and not st.session_state.demo_mode:
                with st.spinner("Analyzing brief quality..."):
                    try:
                        result = checker.check_brief(enriched_text)
                        result["brief_type_checked"] = qc_brief_type
                        result["tier_checked"] = qc_tier
                        st.session_state.quality_result = result
                    except Exception as e:
                        st.error(f"Error: {e}")
            else:
                st.session_state.quality_result = _get_demo_quality_result()

    if st.session_state.quality_result:
        result = st.session_state.quality_result

        # Overall score and letter grade
        st.write("")
        col_score, col_grade, col_radar = st.columns([1, 1, 3])
        with col_score:
            st.markdown("**Overall Score**")
            st.markdown(
                render_quality_score_badge(result["overall_score"]),
                unsafe_allow_html=True,
            )
        with col_grade:
            st.markdown("**Letter Grade**")
            letter = result.get("grade", result.get("letter_grade", ""))
            st.markdown(f"<h1 style='text-align:center; margin-top:10px;'>{letter}</h1>", unsafe_allow_html=True)

        with col_radar:
            # Radar chart
            dims = result.get("dimension_scores", result.get("dimensions", {}))
            categories = list(dims.keys())
            values = list(dims.values())
            # Close the polygon
            categories_closed = categories + [categories[0]]
            values_closed = values + [values[0]]

            fig = go.Figure()
            fig.add_trace(go.Scatterpolar(
                r=values_closed,
                theta=categories_closed,
                fill="toself",
                fillcolor=f"rgba(0, 100, 210, 0.2)",
                line=dict(color=PRIMARY_BLUE, width=2),
                marker=dict(size=6, color=PRIMARY_BLUE),
                name="Quality Score",
            ))
            fig.update_layout(
                polar=dict(
                    radialaxis=dict(visible=True, range=[0, 10]),
                ),
                showlegend=False,
                height=300,
                margin=dict(l=40, r=40, t=20, b=20),
                paper_bgcolor="rgba(0,0,0,0)",
                plot_bgcolor="rgba(0,0,0,0)",
            )
            st.plotly_chart(fig, use_container_width=True)

        # Mistakes found
        st.write("")
        st.markdown("### Issues Found")
        for mistake in result.get("mistakes_found", result.get("mistakes", [])):
            col_sev, col_issue = st.columns([1, 5])
            with col_sev:
                st.markdown(
                    render_severity_badge(mistake["severity"]),
                    unsafe_allow_html=True,
                )
            with col_issue:
                issue_text = mistake.get('issue', mistake.get('description', mistake.get('mistake', '')))
                fix_text = mistake.get('fix', '')
                st.markdown(f"**{mistake.get('section', mistake.get('mistake', ''))}:** {issue_text}")
                if fix_text:
                    st.caption(f"Fix: {fix_text}")

        # Missing sections
        if result.get("missing_sections"):
            st.write("")
            st.markdown("### Missing Sections")
            for section in result["missing_sections"]:
                st.markdown(f"- :x: {section}")

        # Strengths
        if result.get("strengths"):
            st.write("")
            st.markdown("### Strengths")
            for strength in result["strengths"]:
                st.markdown(f"- :white_check_mark: {strength}")

        # Improvement suggestions
        improvements = result.get("improvement_suggestions", result.get("improvements", []))
        if improvements:
            st.write("")
            st.markdown("### Improvement Suggestions")
            for suggestion in improvements:
                st.write(suggestion)



# ---------------------------------------------------------------------------
# Tab 3: Creative Concepts
# ---------------------------------------------------------------------------
def render_creative_concepts() -> None:
    """Render the Creative Concepts tab."""
    st.subheader("Creative Concepts Generator")
    st.caption("Generate creative concept options from a completed brief.")

    # Source selection
    selected_brief_text = ""
    selected_brief = None

    st.markdown("#### Provide your brief")
    selected_brief_text = st.text_area(
        "Paste brief text",
        height=150,
        placeholder="Paste your campaign brief here...",
        label_visibility="collapsed",
    )

    uploaded_concept_brief = st.file_uploader(
        "Or upload a brief document",
        type=["pdf", "txt", "csv", "md", "docx"],
        key="concept_doc_upload",
        help="Upload a brief as PDF, Word, text, or markdown file.",
    )
    if uploaded_concept_brief:
        try:
            if uploaded_concept_brief.name.endswith(".pdf"):
                try:
                    import PyPDF2
                    reader = PyPDF2.PdfReader(uploaded_concept_brief)
                    extracted = "\n".join(page.extract_text() or "" for page in reader.pages[:20])
                except ImportError:
                    extracted = "[PDF support requires PyPDF2: pip install PyPDF2]"
            elif uploaded_concept_brief.name.endswith(".docx"):
                try:
                    from docx import Document as DocxDocument
                    doc = DocxDocument(uploaded_concept_brief)
                    extracted = "\n".join(p.text for p in doc.paragraphs)
                except ImportError:
                    extracted = "[DOCX support requires python-docx: pip install python-docx]"
            else:
                extracted = uploaded_concept_brief.read().decode("utf-8", errors="ignore")

            if extracted.strip():
                selected_brief_text = extracted
                with st.expander(f"Extracted from {uploaded_concept_brief.name}", expanded=False):
                    st.text(extracted[:2000] + ("..." if len(extracted) > 2000 else ""))
        except Exception as e:
            st.error(f"Error reading file: {e}")

    num_concepts = st.slider("Number of concepts", 1, 5, 3)

    if st.button("Generate Concepts", type="primary", use_container_width=True):
        if not selected_brief_text.strip() and not selected_brief:
            st.warning("Please select or paste a brief first.")
        else:
            engine = _get_creative_engine()
            if engine and not st.session_state.demo_mode:
                with st.spinner("Generating creative concepts..."):
                    try:
                        if selected_brief:
                            concepts = engine.generate_concepts(selected_brief, num_concepts)
                        else:
                            concepts = engine.generate_concepts({"summary": selected_brief_text}, num_concepts)
                        st.session_state.creative_concepts = concepts
                    except Exception as e:
                        st.error(f"Error: {e}")
            else:
                st.session_state.creative_concepts = _get_demo_concepts()[:num_concepts]

    if st.session_state.creative_concepts:
        concepts = st.session_state.creative_concepts
        n_cols = min(len(concepts), 3)
        cols = st.columns(n_cols)

        for i, concept in enumerate(concepts):
            with cols[i % n_cols]:
                st.markdown(f"### {concept.get('concept_name', concept.get('name', f'Concept {i+1}'))}")
                st.markdown(f"**Theme:** {concept.get('theme', 'N/A')}")
                st.markdown(f"**Headline:** *{concept.get('headline', '')}*")
                st.markdown(f"**Tagline:** {concept.get('tagline', '')}")
                st.markdown(f"**Tone:** {concept.get('tone', '')}")

                # Appeal score bar
                est_appeal = concept.get("estimated_appeal", {})
                if isinstance(est_appeal, dict):
                    appeal = est_appeal.get("score", concept.get("appeal_score", 7.0))
                else:
                    appeal = concept.get("appeal_score", 7.0)
                pct = int((float(appeal) / 10) * 100)
                st.markdown(f"**Appeal Score:** {appeal}/10")
                st.markdown(
                    f'<div class="appeal-bar-bg"><div class="appeal-bar-fill" style="width:{pct}%"></div></div>',
                    unsafe_allow_html=True,
                )

                with st.expander("Full Concept Details"):
                    st.markdown("**Sample Copy:**")
                    st.write(concept.get("sample_copy", ""))
                    st.markdown("**Visual Direction:**")
                    st.write(concept.get("visual_direction", ""))

                # Download individual concept
                c_name = concept.get('concept_name', concept.get('name', ''))
                concept_md = (
                    f"# Creative Concept: {c_name}\n\n"
                    f"**Theme:** {concept.get('theme', '')}\n"
                    f"**Headline:** {concept.get('headline', '')}\n"
                    f"**Tagline:** {concept.get('tagline', '')}\n"
                    f"**Tone:** {concept.get('tone', '')}\n"
                    f"**Appeal Score:** {appeal}/10\n\n"
                    f"## Sample Copy\n{concept.get('sample_copy', '')}\n\n"
                    f"## Visual Direction\n{concept.get('visual_direction', '')}\n"
                )
                st.download_button(
                    "Download",
                    data=concept_md,
                    file_name=f"concept_{c_name.replace(' ', '_').lower()}.md",
                    mime="text/markdown",
                    key=f"dl_concept_{i}",
                )

        # Email sequence
        st.write("")
        st.markdown("---")
        if st.button("Generate Email Sequence", use_container_width=True):
            if _get_creative_engine() and not st.session_state.demo_mode:
                with st.spinner("Generating email sequence..."):
                    try:
                        engine = _get_creative_engine()
                        emails = engine.generate_email_sequence(selected_brief or {})
                        st.session_state.email_sequence = emails
                    except Exception as e:
                        st.error(f"Error: {e}")
            else:
                st.session_state.email_sequence = _get_demo_email_sequence()

        if st.session_state.email_sequence:
            st.markdown("### 5-Email Drip Campaign")
            for email in st.session_state.email_sequence:
                e_num = email.get("email_number", email.get("email_num", ""))
                subject_raw = email.get("subject_line", email.get("subject", ""))
                if isinstance(subject_raw, dict):
                    e_subject = subject_raw.get("primary", str(subject_raw))
                    e_alts = subject_raw.get("alternatives", [])
                else:
                    e_subject = str(subject_raw)
                    e_alts = []
                e_day = email.get("send_day", email.get("day", ""))
                e_body = email.get("body_copy", email.get("preview", ""))
                e_purpose = email.get("purpose", "")
                e_cta = email.get("cta_text", "")
                e_preview_text = email.get("preview_text", "")

                with st.expander(f"Email {e_num}: {e_subject} ({e_day})"):
                    if e_purpose:
                        st.markdown(f"**Purpose:** {e_purpose}")
                    st.markdown(f"**Subject Line:** {e_subject}")
                    if e_alts:
                        st.caption(f"Alternatives: {' | '.join(str(a) for a in e_alts)}")
                    if e_preview_text:
                        st.markdown(f"**Preview Text:** {e_preview_text}")
                    if e_cta:
                        st.markdown(f"**CTA:** {e_cta}")
                    st.markdown("---")
                    st.markdown(str(e_body))

            # Download email sequence as Word
            try:
                from docx import Document as DocxDocument

                email_doc = DocxDocument()
                email_doc.add_heading("Email Drip Campaign", level=0)

                for email in st.session_state.email_sequence:
                    e_num = email.get("email_number", email.get("email_num", ""))
                    subject_raw = email.get("subject_line", email.get("subject", ""))
                    if isinstance(subject_raw, dict):
                        e_subject = subject_raw.get("primary", str(subject_raw))
                        e_alts = subject_raw.get("alternatives", [])
                    else:
                        e_subject = str(subject_raw)
                        e_alts = []
                    e_day = email.get("send_day", email.get("day", ""))
                    e_body = email.get("body_copy", email.get("preview", ""))
                    e_purpose = email.get("purpose", "")
                    e_cta = email.get("cta_text", "")

                    email_doc.add_heading(f"Email {e_num} — {e_day}", level=2)
                    p = email_doc.add_paragraph()
                    p.add_run("Subject Line: ").bold = True
                    p.add_run(e_subject)
                    if e_alts:
                        p2 = email_doc.add_paragraph()
                        p2.add_run("Alternative Subjects: ").bold = True
                        p2.add_run(" | ".join(str(a) for a in e_alts))
                    if e_purpose:
                        p3 = email_doc.add_paragraph()
                        p3.add_run("Purpose: ").bold = True
                        p3.add_run(str(e_purpose))
                    if e_cta:
                        p4 = email_doc.add_paragraph()
                        p4.add_run("CTA: ").bold = True
                        p4.add_run(str(e_cta))
                    email_doc.add_paragraph("")
                    email_doc.add_paragraph(str(e_body))
                    email_doc.add_paragraph("")

                _email_buffer = io.BytesIO()
                email_doc.save(_email_buffer)
                st.download_button(
                    "Download Email Sequence as Word (.docx)",
                    data=_email_buffer.getvalue(),
                    file_name="email_drip_campaign.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                )
            except ImportError:
                pass


# ---------------------------------------------------------------------------
# Tab 4: Brief Library
# ---------------------------------------------------------------------------
def render_brief_library() -> None:
    """Render the Brief Library tab."""
    st.subheader("Brief Library")

    # Search and filters
    col_search, col_type, col_tier = st.columns([2, 1, 1])
    with col_search:
        search_query = st.text_input("Search briefs", placeholder="Search by campaign name...")
    with col_type:
        filter_type = st.selectbox("Brief Type", ["All"] + BRIEF_TYPES)
    with col_tier:
        filter_tier = st.selectbox("Launch Tier", ["All"] + list(TIER_DESCRIPTIONS.keys()))

    # Filter briefs
    filtered = st.session_state.briefs
    if search_query:
        filtered = [
            b for b in filtered
            if search_query.lower() in b.get("campaign_name", "").lower()
        ]
    if filter_type != "All":
        filtered = [b for b in filtered if b.get("brief_type") == filter_type]
    if filter_tier != "All":
        filtered = [b for b in filtered if b.get("launch_tier") == filter_tier]

    if not filtered:
        st.info("No briefs found matching your criteria.")
    else:
        st.caption(f"Showing {len(filtered)} brief(s)")

        for i, brief in enumerate(filtered):
            with st.container():
                col_info, col_actions = st.columns([3, 2])
                with col_info:
                    tier_html = render_tier_badge(brief.get("launch_tier", "Tier 2"))
                    score = brief.get("quality_score", 0)
                    score_color = ACCENT_GREEN if score > 7 else (ACCENT_YELLOW if score >= 5 else ACCENT_RED)
                    st.markdown(
                        f"### {brief.get('campaign_name', 'Untitled')} "
                        f"{tier_html}",
                        unsafe_allow_html=True,
                    )
                    date_str = brief.get("created", "")[:10]
                    st.caption(
                        f"{brief.get('brief_type', '')}  |  "
                        f"{date_str}  |  "
                        f"Quality: **{score:.1f}**/10"
                    )

                with col_actions:
                    btn_cols = st.columns(4)
                    with btn_cols[0]:
                        if st.button("Edit", key=f"edit_{i}", use_container_width=True):
                            st.session_state.current_brief = brief.copy()
                            st.session_state.wizard_step = 1
                            st.info("Brief loaded into Builder. Switch to the AI Brief Builder tab to edit.")
                    with btn_cols[1]:
                        if st.button("Duplicate", key=f"dup_{i}", use_container_width=True):
                            dup = brief.copy()
                            dup["id"] = str(uuid.uuid4())
                            dup["campaign_name"] = f"{brief['campaign_name']} (Copy)"
                            dup["created"] = datetime.now().isoformat()
                            st.session_state.briefs.append(dup)
                            st.rerun()
                    with btn_cols[2]:
                        if st.button("Delete", key=f"del_{i}", use_container_width=True):
                            st.session_state.briefs = [
                                b for b in st.session_state.briefs
                                if b.get("id") != brief.get("id")
                            ]
                            st.rerun()
                    with btn_cols[3]:
                        md = brief_to_markdown(brief)
                        st.download_button(
                            "Download",
                            data=md,
                            file_name=f"{brief.get('campaign_name', 'brief').replace(' ', '_').lower()}.md",
                            mime="text/markdown",
                            key=f"dl_{i}",
                            use_container_width=True,
                        )

                with st.expander("View Full Brief"):
                    st.markdown(brief_to_markdown(brief))

                st.divider()

    # Export all
    if st.session_state.briefs:
        st.write("")
        all_md = "\n\n---\n\n".join(brief_to_markdown(b) for b in st.session_state.briefs)
        st.download_button(
            "Export All Briefs",
            data=all_md,
            file_name="all_campaign_briefs.md",
            mime="text/markdown",
            use_container_width=True,
        )


# ---------------------------------------------------------------------------
# Tab 5: Templates
# ---------------------------------------------------------------------------
TEMPLATES = [
    {
        "icon": "🚀",
        "name": "Product Launch",
        "description": "Full campaign brief for a major product or feature launch.",
        "tier": "Tier 1",
        "defaults": {
            "brief_type": "Campaign Brief",
            "launch_tier": "Tier 1",
            "background": "A new product/feature is launching that requires full marketing support across all channels.",
            "objective": "Drive [X]% adoption among [target segment] within [timeframe] of launch.",
        },
    },
    {
        "icon": "⬆️",
        "name": "Feature Update",
        "description": "Light brief for an enhancement to an existing product.",
        "tier": "Tier 2",
        "defaults": {
            "brief_type": "Campaign Brief",
            "launch_tier": "Tier 2",
            "background": "An existing feature is being improved with significant enhancements that users should know about.",
            "objective": "Increase awareness of [feature update] to [X]% of active users within [timeframe].",
        },
    },
    {
        "icon": "⚔️",
        "name": "Competitive Response",
        "description": "Rapid-response brief when a competitor makes a major move.",
        "tier": "Tier 1",
        "defaults": {
            "brief_type": "Campaign Brief",
            "launch_tier": "Tier 1",
            "background": "A key competitor has launched [feature/product] that directly challenges our position in [area].",
            "objective": "Counter competitive narrative and reinforce our leadership position within [timeframe].",
        },
    },
    {
        "icon": "🎄",
        "name": "Seasonal Campaign",
        "description": "Time-bound campaign aligned with a holiday or seasonal event.",
        "tier": "Tier 2",
        "defaults": {
            "brief_type": "Campaign Brief",
            "launch_tier": "Tier 2",
            "background": "[Season/Holiday] is approaching and we need to drive engagement with timely, relevant messaging.",
            "objective": "Increase [metric] by [X]% during the [season/holiday] period vs. prior year.",
        },
    },
    {
        "icon": "🔄",
        "name": "Re-engagement Campaign",
        "description": "Win back lapsed users or inactive sellers.",
        "tier": "Tier 2",
        "defaults": {
            "brief_type": "Campaign Brief",
            "launch_tier": "Tier 2",
            "background": "[X]% of users who were active [timeframe] ago have become inactive. We need to re-engage them.",
            "objective": "Re-activate [X]% of lapsed [users/sellers] within [timeframe].",
        },
    },
    {
        "icon": "🌟",
        "name": "Brand Awareness",
        "description": "Top-of-funnel campaign to build brand recognition.",
        "tier": "Tier 0",
        "defaults": {
            "brief_type": "Creative Brief",
            "launch_tier": "Tier 0",
            "background": "Brand awareness in [market/segment] is below target. We need a high-impact campaign to establish presence.",
            "objective": "Increase unaided brand awareness from [X]% to [Y]% in [market] within [timeframe].",
        },
    },
]


def render_templates() -> None:
    """Render the Templates tab with a 2x3 grid."""
    st.subheader("Campaign Brief Templates")
    st.caption("Select a template to pre-fill the Brief Builder with relevant defaults.")

    # 2x3 grid
    for row_start in range(0, len(TEMPLATES), 3):
        cols = st.columns(3)
        for col_idx, template_idx in enumerate(range(row_start, min(row_start + 3, len(TEMPLATES)))):
            template = TEMPLATES[template_idx]
            with cols[col_idx]:
                tier_html = render_tier_badge(template["tier"])
                st.markdown(
                    f"<div class='template-card'>"
                    f"<div style='font-size:2.5em;'>{template['icon']}</div>"
                    f"<h4>{template['name']}</h4>"
                    f"<p style='color:#666; font-size:0.9em;'>{template['description']}</p>"
                    f"{tier_html}"
                    f"</div>",
                    unsafe_allow_html=True,
                )

                if st.button(
                    "Use Template",
                    key=f"template_{template_idx}",
                    use_container_width=True,
                ):
                    st.session_state.current_brief = template["defaults"].copy()
                    st.session_state.current_brief["campaign_name"] = ""
                    st.session_state.wizard_step = 1
                    # Reset generated content
                    st.session_state.generated_names = []
                    st.session_state.ai_insight = None
                    st.session_state.ai_positioning = None
                    st.session_state.ai_messages = None
                    st.session_state.ai_smp = None
                    st.session_state.ai_channels = None
                    st.session_state.ai_deliverables = None
                    st.session_state.ai_timeline = None
                    st.session_state.generated_kpis = None
                    st.session_state.generated_raci = None
                    st.session_state.audience_profile = None
                    st.session_state.smart_objective = None
                    st.session_state.budget_breakdown = None
                    st.session_state.ai_background = None
                    st.session_state.ai_generated_objective = None
                    st.session_state.ai_generated_audience = None
                    st.success(f"Template '{template['name']}' loaded! Switch to the AI Brief Builder tab.")

                with st.expander("Preview Template"):
                    for key, value in template["defaults"].items():
                        label = key.replace("_", " ").title()
                        st.markdown(f"**{label}:** {value}")


# ---------------------------------------------------------------------------
# Main app
# ---------------------------------------------------------------------------
def main() -> None:
    """Main application entry point."""
    init_session_state()

    # Sidebar branding
    with st.sidebar:
        st.markdown("""
        <div style="text-align: left; padding: 0.5rem 0 1.5rem; border-bottom: 1px solid rgba(255,255,255,0.08); margin-bottom: 1rem;">
            <div style="display: flex; align-items: center; gap: 12px;">
                <div style="width: 32px; height: 32px; background: #0064D2; border-radius: 8px; display: flex; align-items: center; justify-content: center;">
                    <span style="color: white; font-weight: 800; font-size: 14px;">CB</span>
                </div>
                <span style="font-size: 1.1rem; font-weight: 700; color: #FFFFFF; letter-spacing: -0.02em;">Campaign Brief</span>
            </div>
            <p style="font-size: 0.7rem; color: #6C8EAD; margin-top: 6px; margin-left: 44px;">v1.0.0</p>
        </div>
        """, unsafe_allow_html=True)

    # Header
    st.markdown("""
    <div style="margin-bottom: 2rem;">
        <h1 style="font-size: 2rem; font-weight: 800; color: #141d23; letter-spacing: -0.02em; margin-bottom: 4px;">
            Campaign Brief Generator
        </h1>
        <p style="color: #727785; font-size: 0.9rem;">
            AI-powered campaign brief creation with industry best practices
        </p>
    </div>
    """, unsafe_allow_html=True)

    # Demo mode banner
    if st.session_state.demo_mode:
        st.info(
            "Running in **Demo Mode**. All features work with sample data. "
            "Add `ANTHROPIC_API_KEY` to `.env` for live AI generation."
        )

    # Tabs
    tab1, tab2, tab3 = st.tabs([
        "AI Brief Builder",
        "Brief Quality Checker",
        "Creative Concepts",
    ])

    with tab1:
        render_brief_builder()

    with tab2:
        render_quality_checker()

    with tab3:
        render_creative_concepts()


if __name__ == "__main__":
    main()
